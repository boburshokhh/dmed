"""Замена плейсхолдеров в DOCX документах"""
import re
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import DOCX_FONT_NAME


def replace_placeholders_in_document(doc, replacements):
    """Заменяет плейсхолдеры во всем документе"""
    pin_code_first_occurrence = {'found': False}
    
    # Заменяем плейсхолдеры в параграфах
    for paragraph in doc.paragraphs:
        replace_placeholders_with_font(paragraph, replacements, pin_code_first_occurrence)
    
    # Заменяем плейсхолдеры в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_with_font(paragraph, replacements, pin_code_first_occurrence)
    
    # Заменяем плейсхолдеры во всех типах колонтитулов
    for section in doc.sections:
        process_header_footer(section.header, replacements, pin_code_first_occurrence, "header")
        process_header_footer(section.footer, replacements, pin_code_first_occurrence, "footer")
        
        # Колонтитулы первой страницы
        try:
            if hasattr(section, 'first_page_header') and section.first_page_header:
                process_header_footer(section.first_page_header, replacements, pin_code_first_occurrence, "first_page_header")
        except:
            pass
        
        try:
            if hasattr(section, 'first_page_footer') and section.first_page_footer:
                process_header_footer(section.first_page_footer, replacements, pin_code_first_occurrence, "first_page_footer")
        except:
            pass
        
        # Колонтитулы четных/нечетных страниц
        try:
            if hasattr(section, 'even_page_header') and section.even_page_header:
                process_header_footer(section.even_page_header, replacements, pin_code_first_occurrence, "even_page_header")
        except:
            pass
        
        try:
            if hasattr(section, 'even_page_footer') and section.even_page_footer:
                process_header_footer(section.even_page_footer, replacements, pin_code_first_occurrence, "even_page_footer")
        except:
            pass
        
        try:
            if hasattr(section, 'odd_page_header') and section.odd_page_header:
                process_header_footer(section.odd_page_header, replacements, pin_code_first_occurrence, "odd_page_header")
        except:
            pass
        
        try:
            if hasattr(section, 'odd_page_footer') and section.odd_page_footer:
                process_header_footer(section.odd_page_footer, replacements, pin_code_first_occurrence, "odd_page_footer")
        except:
            pass


def process_header_footer(header_footer, replacements, pin_code_first_occurrence, header_footer_name=""):
    """Обрабатывает все параграфы и таблицы в колонтитуле"""
    if not header_footer:
        return
    
    try:
        # Обрабатываем параграфы
        for paragraph in header_footer.paragraphs:
            try:
                replace_placeholders_with_font(paragraph, replacements, pin_code_first_occurrence)
            except:
                pass
        
        # Обрабатываем таблицы
        for table in header_footer.tables:
            try:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            try:
                                replace_placeholders_with_font(paragraph, replacements, pin_code_first_occurrence)
                            except:
                                pass
            except:
                pass
    except:
        pass


def replace_placeholders_with_font(paragraph, replacements, pin_code_first_occurrence, font_name=DOCX_FONT_NAME):
    """Заменяет плейсхолдеры в параграфе, устанавливая шрифт Arial для замененных значений."""
    try:
        # Получаем полный текст параграфа
        full_text = paragraph.text
        runs_text = []
        for run in paragraph.runs:
            runs_text.append(run.text)
        combined_runs_text = ''.join(runs_text)
        check_text = full_text if len(full_text) >= len(combined_runs_text) else combined_runs_text
        
        # Проверяем, есть ли плейсхолдеры
        has_placeholder = False
        for key in replacements.keys():
            if key in check_text:
                has_placeholder = True
                break
        
        if not has_placeholder:
            return
        
        # Сохраняем выравнивание и стиль параграфа
        alignment = paragraph.alignment
        original_font_size = None
        original_italic = None
        
        if paragraph.runs:
            try:
                first_run = paragraph.runs[0]
                original_font_size = first_run.font.size
                original_italic = first_run.font.italic
            except:
                pass
        
        # Очищаем параграф
        paragraph.clear()
        if alignment:
            paragraph.alignment = alignment
        
        # Разбиваем текст на части и обрабатываем каждый плейсхолдер отдельно
        text_to_process = check_text if check_text else full_text
        pattern = '|'.join(re.escape(key) for key in replacements.keys())
        parts = re.split(f'({pattern})', text_to_process)
        
        for part in parts:
            if not part:
                continue
            
            if part in replacements:
                value = str(replacements[part])
                run = paragraph.add_run(value)
                
                # Специальная обработка для {{pin_code}}
                if part == '{{pin_code}}':
                    pin_size = Pt(14)
                    if not pin_code_first_occurrence['found']:
                        pin_code_first_occurrence['found'] = True
                    
                    try:
                        run.font.size = pin_size
                    except:
                        try:
                            if run._element.rPr is None:
                                run._element.get_or_add_rPr()
                            sz_value = int(pin_size.pt * 2)
                            sz = OxmlElement('w:sz')
                            sz.set(qn('w:val'), str(sz_value))
                            run._element.rPr.append(sz)
                        except:
                            pass
                    
                    # Устанавливаем шрифт
                    for font_name_linux in ['Arial', 'DejaVu Sans', 'Liberation Sans', 'Calibri', 'Helvetica']:
                        try:
                            run.font.name = font_name_linux
                            break
                        except:
                            continue
                    
                    # Устанавливаем жирный шрифт
                    try:
                        run.bold = True
                    except:
                        try:
                            if run._element.rPr is None:
                                run._element.get_or_add_rPr()
                            b = run._element.rPr.get_or_add_b()
                            b.set(qn('w:val'), 'true')
                        except:
                            pass
                    
                    # Дополнительная проверка жирности через XML
                    try:
                        if run._element.rPr is None:
                            run._element.get_or_add_rPr()
                        b_elements = run._element.rPr.findall(qn('w:b'))
                        if not b_elements:
                            b = OxmlElement('w:b')
                            b.set(qn('w:val'), 'true')
                            run._element.rPr.append(b)
                        else:
                            for b in b_elements:
                                b.set(qn('w:val'), 'true')
                    except:
                        pass
                else:
                    # Значение переменной - применяем стандартное форматирование
                    try:
                        run.font.name = 'Arial'
                    except:
                        try:
                            run.font.name = font_name
                        except:
                            pass
                    if original_font_size:
                        try:
                            run.font.size = original_font_size
                        except:
                            pass
                    
                    # Специальная обработка для {{issue_date}} и {{doc_number}}
                    if part == '{{issue_date}}' or part == '{{doc_number}}':
                        try:
                            run.font.bold = True
                        except:
                            pass
                        try:
                            if run._element.rPr is None:
                                run._element.get_or_add_rPr()
                            b_elements = run._element.rPr.findall(qn('w:b'))
                            if not b_elements:
                                b = OxmlElement('w:b')
                                b.set(qn('w:val'), 'true')
                                run._element.rPr.append(b)
                            else:
                                for b in b_elements:
                                    b.set(qn('w:val'), 'true')
                        except:
                            pass
                    else:
                        try:
                            run.font.bold = False
                        except:
                            pass
            else:
                # Обычный текст (названия переменных) - должны быть жирными
                run = paragraph.add_run(part)
                try:
                    run.font.name = 'Arial'
                except:
                    try:
                        run.font.name = font_name
                    except:
                        pass
                if original_font_size:
                    try:
                        run.font.size = original_font_size
                    except:
                        pass
                try:
                    run.font.bold = True
                except:
                    pass
                try:
                    if run._element.rPr is None:
                        run._element.get_or_add_rPr()
                    b_elements = run._element.rPr.findall(qn('w:b'))
                    if not b_elements:
                        b = OxmlElement('w:b')
                        b.set(qn('w:val'), 'true')
                        run._element.rPr.append(b)
                    else:
                        for b in b_elements:
                            b.set(qn('w:val'), 'true')
                except:
                    pass
                if original_italic is not None:
                    try:
                        run.font.italic = original_italic
                    except:
                        pass
        
    except Exception as e:
        import traceback
        print(f"[ERROR] Ошибка при замене плейсхолдеров: {e}")
        print(traceback.format_exc())
        try:
            simple_text = check_text if 'check_text' in locals() else full_text
            for key, value in replacements.items():
                simple_text = simple_text.replace(key, str(value))
            paragraph.text = simple_text
        except:
            pass

