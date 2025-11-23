"""Обработка QR-кодов и PIN-кодов в DOCX документах"""
import os
import re
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import FRONTEND_URL
from utils import generate_qr_code
from PIL import Image


def add_qr_code_to_docx(doc, pin_code, app=None, document_uuid=None):
    """Добавляет QR-код в DOCX документ с PIN-кодом"""
    try:
        # Генерируем URL для QR-кода
        if document_uuid:
            frontend_url = FRONTEND_URL.rstrip('/')
            qr_url = f"{frontend_url}/access/{document_uuid}"
        else:
            try:
                from flask import url_for
                qr_url = url_for('verify_document', _external=True)
            except RuntimeError:
                qr_url = '/verify'
        
        qr_img = generate_qr_code(qr_url)
        
        if qr_img is None:
            raise Exception("generate_qr_code() вернул None")
        
        if not isinstance(qr_img, Image.Image):
            raise Exception(f"generate_qr_code() вернул не PIL.Image, а {type(qr_img)}")
        
        # Сохраняем QR-код во временный файл
        upload_folder = app.config.get('UPLOAD_FOLDER', 'static/generated_documents') if app else 'static/generated_documents'
        qr_temp_path = os.path.join(upload_folder, f'qr_temp_{pin_code}.png')
        if not os.path.exists(os.path.dirname(qr_temp_path)):
            os.makedirs(os.path.dirname(qr_temp_path), exist_ok=True)
        
        # Конвертируем в RGBA если нужно
        if qr_img.mode != 'RGBA':
            try:
                qr_img = qr_img.convert('RGBA')
            except:
                pass
        
        # Сохраняем изображение
        try:
            with open(qr_temp_path, 'wb') as f:
                qr_img.save(f, 'PNG', optimize=True)
            
            if os.path.exists(qr_temp_path):
                file_size = os.path.getsize(qr_temp_path)
                if file_size == 0:
                    raise Exception(f"Файл QR-кода пустой: {qr_temp_path}")
            else:
                raise Exception(f"Файл QR-кода не был создан: {qr_temp_path}")
        except Exception as save_error:
            print(f"[ERROR] Ошибка при сохранении QR-кода: {save_error}")
            import traceback
            print(traceback.format_exc())
            raise
        
        qr_placeholder_found = False
        
        # Ищем плейсхолдер в параграфах
        for paragraph in doc.paragraphs:
            if replace_qr_placeholder(paragraph, doc, pin_code, qr_temp_path, False, None):
                qr_placeholder_found = True
                break
        
        # Ищем плейсхолдер в таблицах
        if not qr_placeholder_found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if replace_qr_placeholder(paragraph, doc, pin_code, qr_temp_path, True, cell):
                                qr_placeholder_found = True
                                break
                        if qr_placeholder_found:
                            break
                    if qr_placeholder_found:
                        break
                if qr_placeholder_found:
                    break
        
        # Ищем плейсхолдер в колонтитулах
        if not qr_placeholder_found:
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if replace_qr_placeholder(paragraph, doc, pin_code, qr_temp_path, False, None):
                            qr_placeholder_found = True
                            break
                    if not qr_placeholder_found:
                        for table in section.header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if replace_qr_placeholder(paragraph, doc, pin_code, qr_temp_path, True, cell):
                                            qr_placeholder_found = True
                                            break
                                    if qr_placeholder_found:
                                        break
                                if qr_placeholder_found:
                                    break
                            if qr_placeholder_found:
                                break
                
                if qr_placeholder_found:
                    break
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if replace_qr_placeholder(paragraph, doc, pin_code, qr_temp_path, False, None):
                            qr_placeholder_found = True
                            break
                    if not qr_placeholder_found:
                        for table in section.footer.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if replace_qr_placeholder(paragraph, doc, pin_code, qr_temp_path, True, cell):
                                            qr_placeholder_found = True
                                            break
                                    if qr_placeholder_found:
                                        break
                                if qr_placeholder_found:
                                    break
                            if qr_placeholder_found:
                                break
                
                if qr_placeholder_found:
                    break
        
        # Если плейсхолдер не найден, добавляем таблицу в конец документа
        if not qr_placeholder_found:
            create_pin_qr_table(doc, pin_code, qr_temp_path, 'right', 'left')
        
        # Удаляем временный файл
        if os.path.exists(qr_temp_path):
            os.remove(qr_temp_path)
            
    except Exception as e:
        print(f"[ERROR] Ошибка при добавлении QR-кода: {e}")
        import traceback
        print(traceback.format_exc())


def parse_placeholder_alignment(text):
    """
    Парсит плейсхолдер и определяет выравнивание из него.
    Возвращает: (placeholder_type, pin_alignment, qr_alignment)
    """
    patterns = [
        (r'\{\{qr_code:(left|right)\}\}', 'qr_code'),
        (r'\{\{pin_code_with_qr:(left|right)\}\}', 'pin_code_with_qr'),
        (r'\{\{pin_code:(left|right)\}\}', 'pin_code'),
    ]
    
    for pattern, ptype in patterns:
        match = re.search(pattern, text)
        if match:
            alignment = match.group(1)
            if ptype == 'qr_code':
                return (ptype, 'right', alignment)
            elif ptype == 'pin_code':
                return (ptype, alignment, 'left')
            else:  # pin_code_with_qr
                return (ptype, alignment, alignment)
    
    # Старые форматы
    if '{{qr_code}}' in text or '{{pin_code_with_qr}}' in text:
        return ('pin_code_with_qr', 'right', 'left')
    elif '{{pin_code}}' in text:
        return ('pin_code', 'right', None)
    
    return (None, 'right', 'left')


def replace_qr_placeholder(paragraph, doc, pin_code, qr_temp_path, in_table_cell=False, cell=None):
    """Заменяет плейсхолдер QR-кода на таблицу с PIN-кодом и QR-кодом"""
    placeholder_type, pin_align, qr_align = parse_placeholder_alignment(paragraph.text)
    
    if placeholder_type:
        if in_table_cell and cell:
            # Вложенная таблица в ячейке
            for para in cell.paragraphs[:]:
                p_element = para._element
                p_element.getparent().remove(p_element)
            
            inner_table = cell.add_table(rows=1, cols=2)
            inner_table.style = None
            inner_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            inner_table.columns[0].width = Cm(3.0)
            inner_table.columns[1].width = Cm(5.0)
            
            # Убираем spacing для вложенной таблицы
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tbl_pr = inner_table._element.tblPr
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                inner_table._element.insert(0, tbl_pr)
            tbl_cell_spacing = OxmlElement('w:tblCellSpacing')
            tbl_cell_spacing.set(qn('w:w'), '0')
            tbl_cell_spacing.set(qn('w:type'), 'dxa')
            old_cell_spacing = tbl_pr.find(qn('w:tblCellSpacing'))
            if old_cell_spacing is not None:
                tbl_pr.remove(old_cell_spacing)
            tbl_pr.append(tbl_cell_spacing)
            
            setup_pin_cell(inner_table.rows[0].cells[0], pin_code, pin_align)
            setup_qr_cell(inner_table.rows[0].cells[1], qr_temp_path, qr_align if qr_align else 'left')
        else:
            # Обычная таблица
            parent = paragraph._element.getparent()
            qr_align_final = qr_align if qr_align is not None else 'left'
            table = create_pin_qr_table(doc, pin_code, qr_temp_path, pin_align, qr_align_final)
            table_element = table._element
            parent.insert(parent.index(paragraph._element), table_element)
            parent.remove(paragraph._element)
        
        return True
    return False


def create_pin_qr_table(doc, pin_code, qr_temp_path, pin_alignment='right', qr_alignment='left'):
    """Создает таблицу с PIN-кодом и QR-кодом"""
    table = doc.add_table(rows=1, cols=2)
    table.style = None
    
    # Убираем spacing между ячейками
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl_pr = table._element.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    tbl_cell_spacing = OxmlElement('w:tblCellSpacing')
    tbl_cell_spacing.set(qn('w:w'), '0')
    tbl_cell_spacing.set(qn('w:type'), 'dxa')
    old_cell_spacing = tbl_pr.find(qn('w:tblCellSpacing'))
    if old_cell_spacing is not None:
        tbl_pr.remove(old_cell_spacing)
    tbl_pr.append(tbl_cell_spacing)
    
    # Настраиваем ширину колонок
    table.columns[0].width = Cm(3.0)
    table.columns[1].width = Cm(5.0)
    
    # Настраиваем ячейки
    setup_pin_cell(table.rows[0].cells[0], pin_code, pin_alignment)
    setup_qr_cell(table.rows[0].cells[1], qr_temp_path, qr_alignment)
    
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    return table


def setup_pin_cell(cell, pin_code, alignment):
    """Настраивает ячейку с PIN-кодом"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Убираем отступы через XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._element.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        cell._element.append(tc_pr)
    
    no_wrap = OxmlElement('w:noWrap')
    tc_pr.append(no_wrap)
    
    tc_mar = OxmlElement('w:tcMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        margin_elem = OxmlElement(f'w:{margin}')
        margin_elem.set(qn('w:w'), '0')
        margin_elem.set(qn('w:type'), 'dxa')
        tc_mar.append(margin_elem)
    
    old_tc_mar = tc_pr.find(qn('w:tcMar'))
    if old_tc_mar is not None:
        tc_pr.remove(old_tc_mar)
    tc_pr.append(tc_mar)
    
    para_pin = cell.paragraphs[0]
    para_pin.alignment = WD_ALIGN_PARAGRAPH.RIGHT if alignment == 'right' else WD_ALIGN_PARAGRAPH.LEFT
    para_pin_format = para_pin.paragraph_format
    para_pin_format.space_before = Pt(0)
    para_pin_format.space_after = Pt(0)
    para_pin_format.left_indent = Pt(35)
    para_pin_format.right_indent = Pt(0)
    para_pin_format.widow_control = False
    para_pin_format.keep_together = True
    
    para_pin.clear()
    pin_text = str(pin_code).strip().replace('\n', '').replace('\r', '').replace(' ', '')
    run_pin = para_pin.add_run(pin_text)
    run_pin.font.size = Pt(20)
    run_pin.font.bold = True
    
    r_pr = run_pin._element.get_or_add_rPr()
    no_break = OxmlElement('w:noBreak')
    r_pr.append(no_break)
    
    for font_name in ['Arial', 'DejaVu Sans', 'Liberation Sans', 'Calibri', 'Helvetica']:
        try:
            run_pin.font.name = font_name
            break
        except:
            continue


def setup_qr_cell(cell, qr_temp_path, alignment):
    """Настраивает ячейку с QR-кодом"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Убираем отступы через XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._element.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        cell._element.append(tc_pr)
    
    tc_mar = OxmlElement('w:tcMar')
    for margin in ['top', 'bottom', 'right', 'left']:
        margin_elem = OxmlElement(f'w:{margin}')
        margin_elem.set(qn('w:w'), '0')
        margin_elem.set(qn('w:type'), 'dxa')
        tc_mar.append(margin_elem)
    
    old_tc_mar = tc_pr.find(qn('w:tcMar'))
    if old_tc_mar is not None:
        tc_pr.remove(old_tc_mar)
    tc_pr.append(tc_mar)
    
    para_qr = cell.paragraphs[0]
    para_qr.alignment = WD_ALIGN_PARAGRAPH.LEFT if alignment == 'left' else WD_ALIGN_PARAGRAPH.RIGHT
    para_qr_format = para_qr.paragraph_format
    para_qr_format.space_before = Pt(0)
    para_qr_format.space_after = Pt(0)
    para_qr_format.left_indent = Pt(10)
    para_qr_format.right_indent = Pt(0)
    
    run_qr = para_qr.add_run()
    qr_width_inches = 1.21
    run_qr.add_picture(qr_temp_path, width=Inches(qr_width_inches))

