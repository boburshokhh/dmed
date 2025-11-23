"""Создание стандартных DOCX шаблонов"""
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import DOCX_FONT_NAME
from docx_formatter import setup_a4_page_format


def create_default_docx_template():
    """Создает стандартный DOCX шаблон если его нет"""
    doc = Document()
    setup_a4_page_format(doc)
    
    try:
        title = doc.add_heading("Ta'lim olayotgan shaxslar uchun mehnatga layoqatsizlik ma'lumotnomasi", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        title = doc.add_paragraph("Ta'lim olayotgan shaxslar uchun mehnatga layoqatsizlik ma'lumotnomasi")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].font.bold = True
            title.runs[0].font.name = DOCX_FONT_NAME
    
    org_para = doc.add_paragraph("O'zbekiston Respublikasi Sog'liqni saqlash vazirligi")
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_para = doc.add_paragraph("4 - sonli oilaviy poliklinika")
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Ro'yhatga olingan sana: {{issue_date}}")
    doc.add_paragraph(f"No {{doc_number}}")
    
    try:
        doc.add_heading('Vaqtincha mehnatga layoqatsiz fuqaro haqidagi ma\'lumotlar:', 1)
    except:
        para = doc.add_paragraph('Vaqtincha mehnatga layoqatsiz fuqaro haqidagi ma\'lumotlar:')
        para.runs[0].font.bold = True
        para.runs[0].font.name = DOCX_FONT_NAME
    
    doc.add_paragraph(f"FISH: {{patient_name}}")
    doc.add_paragraph(f"Jinsi: {{gender}}")
    doc.add_paragraph(f"JShShIR: {{jshshir}}")
    doc.add_paragraph(f"Yoshi: {{age}}")
    doc.add_paragraph(f"Yashash manzili: {{address}}")
    doc.add_paragraph(f"Biriktirilgan tibbiy muassasa: {{attached_medical_institution}}")
    doc.add_paragraph(f"Tashxis (KXT-10 kodi va Nomi): {{diagnosis_icd10_code}}: {{diagnosis}}")
    doc.add_paragraph(f"Yakuniy tashxis (Nomi va KXT-10 kodi): {{final_diagnosis_icd10_code}}: {{final_diagnosis}}")
    doc.add_paragraph(f"Ro'yhatga olingan sana: {{issue_date}}")
    doc.add_paragraph(f"No: {{doc_number}}")
    doc.add_paragraph(f"Davolovchi shifokor FISH: {{doctor_name}}")
    doc.add_paragraph(f"Bo'lim boshlig'i (mas'ul shaxs) FISH: {{department_head_name}}")
    doc.add_paragraph(f"Ishdan ozod etilgan kunlar: {{days_off_period}}")
    
    try:
        doc.add_heading('Shifokor ma\'lumotlari:', 1)
    except:
        para = doc.add_paragraph('Shifokor ma\'lumotlari:')
        if para.runs:
            para.runs[0].font.bold = True
            para.runs[0].font.name = DOCX_FONT_NAME
        else:
            run = para.add_run('Shifokor ma\'lumotlari:')
            run.font.bold = True
            run.font.name = DOCX_FONT_NAME
    
    doc.add_paragraph(f"FISH: {{doctor_name}}")
    doc.add_paragraph(f"Lavozim: {{doctor_position}}")
    doc.add_paragraph(f"PIN-kod: {{pin_code}}")
    
    return doc

