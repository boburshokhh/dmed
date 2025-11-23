"""Генерация DOCX документов"""
import os
import uuid
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import DOCX_FONT_NAME
from storage import storage_manager
from utils import generate_qr_code
from PIL import Image


def fill_docx_template(document_data, template_path=None, app=None):
    """Заполняет DOCX шаблон данными из формы"""
    try:
        # Определяем путь к шаблону
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            # Используем стандартный шаблон или создаем новый
            template_folder = app.config.get('TEMPLATE_FOLDER', 'templates/docx_templates') if app else 'templates/docx_templates'
            default_template = os.path.join(template_folder, 'template.docx')
            if os.path.exists(default_template):
                doc = Document(default_template)
            else:
                # Создаем новый документ если шаблона нет
                doc = create_default_docx_template()
        
        # Настраиваем размеры страницы A4 и минимальные отступы для всех документов
        from docx.shared import Cm
        sections = doc.sections
        for section in sections:
            section.page_height = Cm(29.7)  # A4 высота
            section.page_width = Cm(21.0)   # A4 ширина
            section.left_margin = Cm(1.0)   # Минимальные отступы
            section.right_margin = Cm(1.0)
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
        
        # Подготавливаем данные для замены
        replacements = {
            '{{doc_number}}': document_data.get('doc_number', ''),
            '{{pin_code}}': document_data.get('pin_code', ''),
            '{{uuid}}': document_data.get('uuid', ''),
            '{{patient_name}}': document_data.get('patient_name', ''),
            '{{gender}}': document_data.get('gender', ''),
            '{{age}}': document_data.get('age', ''),
            '{{jshshir}}': document_data.get('jshshir', ''),
            '{{address}}': document_data.get('address', ''),
            '{{attached_medical_institution}}': document_data.get('attached_medical_institution', ''),
            '{{diagnosis}}': document_data.get('diagnosis', ''),
            '{{diagnosis_icd10_code}}': document_data.get('diagnosis_icd10_code', ''),
            '{{final_diagnosis}}': document_data.get('final_diagnosis', ''),
            '{{final_diagnosis_icd10_code}}': document_data.get('final_diagnosis_icd10_code', ''),
            '{{organization}}': document_data.get('organization', ''),
            '{{doctor_name}}': document_data.get('doctor_name', ''),
            '{{doctor_position}}': document_data.get('doctor_position', ''),
            '{{department_head_name}}': document_data.get('department_head_name', ''),
        }
        
        # Форматируем даты освобождения
        days_off_from = document_data.get('days_off_from')
        days_off_to = document_data.get('days_off_to')
        
        if days_off_from:
            if isinstance(days_off_from, str):
                try:
                    dt = datetime.fromisoformat(days_off_from.replace('Z', '+00:00'))
                    replacements['{{days_off_from}}'] = dt.strftime("%d.%m.%Y")
                except:
                    replacements['{{days_off_from}}'] = days_off_from[:10] if len(days_off_from) >= 10 else days_off_from
            else:
                replacements['{{days_off_from}}'] = days_off_from.strftime("%d.%m.%Y")
        else:
            replacements['{{days_off_from}}'] = ''
        
        if days_off_to:
            if isinstance(days_off_to, str):
                try:
                    dt = datetime.fromisoformat(days_off_to.replace('Z', '+00:00'))
                    replacements['{{days_off_to}}'] = dt.strftime("%d.%m.%Y")
                except:
                    replacements['{{days_off_to}}'] = days_off_to[:10] if len(days_off_to) >= 10 else days_off_to
            else:
                replacements['{{days_off_to}}'] = days_off_to.strftime("%d.%m.%Y")
        else:
            replacements['{{days_off_to}}'] = ''
        
        # Формируем строку периода освобождения
        if replacements['{{days_off_from}}'] and replacements['{{days_off_to}}']:
            replacements['{{days_off_period}}'] = f"{replacements['{{days_off_from}}']} - {replacements['{{days_off_to}}']}"
        elif replacements['{{days_off_from}}']:
            replacements['{{days_off_period}}'] = replacements['{{days_off_from}}']
        else:
            replacements['{{days_off_period}}'] = ''
        
        # Форматируем дату выдачи документа (с временем)
        issue_date = document_data.get('issue_date')
        if issue_date:
            if isinstance(issue_date, str):
                try:
                    dt = datetime.fromisoformat(issue_date.replace('Z', '+00:00'))
                    date_str = dt.strftime("%d.%m.%Y %H:%M")
                except:
                    try:
                        dt = datetime.strptime(issue_date[:16], "%Y-%m-%dT%H:%M")
                        date_str = dt.strftime("%d.%m.%Y %H:%M")
                    except:
                        date_str = issue_date[:10] if len(issue_date) >= 10 else issue_date
            else:
                date_str = issue_date.strftime("%d.%m.%Y %H:%M")
        else:
            date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        
        replacements['{{issue_date}}'] = date_str
        replacements['{{current_date}}'] = datetime.now().strftime("%d.%m.%Y %H:%M")
        
        # Флаг для отслеживания первого вхождения {{pin_code}}
        pin_code_first_occurrence = {'found': False}
        
        # Функция для замены плейсхолдеров с сохранением форматирования
        def replace_placeholders_with_font(paragraph, replacements, font_name=DOCX_FONT_NAME):
            """Заменяет плейсхолдеры в параграфе, устанавливая шрифт Arial для замененных значений."""
            try:
                # Получаем полный текст параграфа, объединяя все runs
                full_text = paragraph.text
                
                # Дополнительная проверка: собираем текст из всех runs по отдельности
                # на случай, если переменная разбита на несколько runs
                runs_text = []
                for run in paragraph.runs:
                    runs_text.append(run.text)
                combined_runs_text = ''.join(runs_text)
                
                # Используем более полный текст для проверки
                check_text = full_text if len(full_text) >= len(combined_runs_text) else combined_runs_text
                
                # Проверяем, есть ли плейсхолдеры
                has_placeholder = False
                found_placeholders = []
                for key in replacements.keys():
                    if key in check_text:
                        has_placeholder = True
                        found_placeholders.append(key)
                
                if not has_placeholder:
                    return
                
                print(f"[DEBUG] Найдены плейсхолдеры в параграфе: {found_placeholders}, текст: {check_text[:100]}")
                
                # Сохраняем выравнивание и стиль параграфа
                alignment = paragraph.alignment
                original_font_size = None
                original_bold = None
                original_italic = None
                
                # Проверяем, есть ли жирный текст в параграфе (названия переменных должны быть жирными)
                has_bold_text = False
                if paragraph.runs:
                    try:
                        # Проверяем все runs на наличие жирного текста
                        for run in paragraph.runs:
                            if run.font.bold:
                                has_bold_text = True
                                break
                        first_run = paragraph.runs[0]
                        original_font_size = first_run.font.size
                        original_bold = first_run.font.bold
                        original_italic = first_run.font.italic
                    except:
                        pass
                
                # Если в параграфе есть переменные, названия должны быть жирными
                # Устанавливаем original_bold = True для названий переменных
                if has_placeholder:
                    original_bold = True
                
                # Очищаем параграф
                paragraph.clear()
                if alignment:
                    paragraph.alignment = alignment
                
                # Используем check_text для разбиения (более полный текст)
                text_to_process = check_text if check_text else full_text
                
                # Разбиваем текст на части и обрабатываем каждый плейсхолдер отдельно
                pattern = '|'.join(re.escape(key) for key in replacements.keys())
                parts = re.split(f'({pattern})', text_to_process)
                
                print(f"[DEBUG] Разбит текст на {len(parts)} частей, найдено плейсхолдеров: {len([p for p in parts if p in replacements])}")
                
                for part in parts:
                    if not part:
                        continue
                    
                    if part in replacements:
                        value = str(replacements[part])
                        run = paragraph.add_run(value)
                        
                        # Специальная обработка для {{pin_code}} - ВСЕ вхождения делаем большими и жирными
                        if part == '{{pin_code}}':
                            # Большой PIN-код - применяем ко всем вхождениям
                            # Используем шрифты, доступные на Linux
                            # Сначала устанавливаем размер - это критично
                            pin_size = Pt(14)  # Размер шрифта для PIN-кода (уменьшен на 10 единиц)
                            
                            # Если это первое вхождение, отмечаем его
                            if not pin_code_first_occurrence['found']:
                                pin_code_first_occurrence['found'] = True
                            
                            # Применяем большой размер и жирность ко всем вхождениям
                            try:
                                run.font.size = pin_size
                                print(f"[INFO] Установлен размер шрифта для PIN-кода: {pin_size.pt}pt")
                            except Exception as size_error:
                                print(f"[WARNING] Не удалось установить размер шрифта для PIN-кода: {size_error}")
                                # Пробуем альтернативный способ через XML
                                try:
                                    from docx.oxml.ns import qn
                                    from docx.oxml import OxmlElement
                                    # Убеждаемся, что rPr существует
                                    if run._element.rPr is None:
                                        run._element.get_or_add_rPr()
                                    # Размер в half-points (24pt = 48 half-points)
                                    sz_value = int(pin_size.pt * 2)
                                    sz = OxmlElement('w:sz')
                                    sz.set(qn('w:val'), str(sz_value))
                                    run._element.rPr.append(sz)
                                    print(f"[INFO] Размер установлен через XML: {sz_value} half-points")
                                except Exception as xml_error:
                                    print(f"[WARNING] Альтернативный способ установки размера не сработал: {xml_error}")
                                    pass
                            
                            # Устанавливаем шрифт Arial
                            font_set = False
                            for font_name_linux in ['Arial', 'DejaVu Sans', 'Liberation Sans', 'Calibri', 'Helvetica']:
                                try:
                                    run.font.name = font_name_linux
                                    font_set = True
                                    print(f"[INFO] Установлен шрифт для PIN-кода: {font_name_linux}")
                                    break
                                except Exception as font_error:
                                    continue
                            
                            if not font_set:
                                print("[WARNING] Не удалось установить ни один шрифт для PIN-кода")
                            
                            # Устанавливаем жирный шрифт - применяем все способы
                            bold_set = False
                            try:
                                run.bold = True
                                bold_set = True
                                print("[INFO] Жирный шрифт установлен через run.bold")
                            except Exception as bold_error:
                                print(f"[WARNING] Не удалось установить жирный шрифт через run.bold: {bold_error}")
                            
                            if not bold_set:
                                try:
                                    # Альтернативный способ установки жирного шрифта через XML
                                    from docx.oxml.ns import qn
                                    if run._element.rPr is None:
                                        run._element.get_or_add_rPr()
                                    b = run._element.rPr.get_or_add_b()
                                    b.set(qn('w:val'), 'true')
                                    bold_set = True
                                    print("[INFO] Жирный шрифт установлен через XML")
                                except Exception as xml_bold_error:
                                    print(f"[WARNING] Не удалось установить жирный шрифт через XML: {xml_bold_error}")
                            
                            # Дополнительно: устанавливаем font-weight через прямое обращение к XML
                            try:
                                from docx.oxml.ns import qn
                                from docx.oxml import OxmlElement
                                if run._element.rPr is None:
                                    run._element.get_or_add_rPr()
                                # Убеждаемся, что жирность установлена
                                b_elements = run._element.rPr.findall(qn('w:b'))
                                if not b_elements:
                                    b = OxmlElement('w:b')
                                    b.set(qn('w:val'), 'true')
                                    run._element.rPr.append(b)
                                else:
                                    for b in b_elements:
                                        b.set(qn('w:val'), 'true')
                                print("[INFO] Дополнительная проверка жирности выполнена")
                            except Exception as extra_bold_error:
                                print(f"[WARNING] Дополнительная установка жирности не сработала: {extra_bold_error}")
                        else:
                            # Значение переменной - применяем стандартное форматирование
                            try:
                                run.font.name = 'Arial'
                            except:
                                try:
                                    run.font.name = font_name
                                except:
                                    pass
                            if original_font_size:
                                try:
                                    run.font.size = original_font_size
                                except:
                                    pass
                            
                            # Специальная обработка для {{issue_date}} и {{doc_number}} - делаем жирными
                            if part == '{{issue_date}}' or part == '{{doc_number}}':
                                try:
                                    run.font.bold = True
                                except:
                                    pass
                                # Дополнительно через XML для гарантии
                                try:
                                    from docx.oxml.ns import qn
                                    from docx.oxml import OxmlElement
                                    if run._element.rPr is None:
                                        run._element.get_or_add_rPr()
                                    # Убеждаемся, что жирность установлена
                                    b_elements = run._element.rPr.findall(qn('w:b'))
                                    if not b_elements:
                                        b = OxmlElement('w:b')
                                        b.set(qn('w:val'), 'true')
                                        run._element.rPr.append(b)
                                    else:
                                        for b in b_elements:
                                            b.set(qn('w:val'), 'true')
                                except Exception as xml_bold_error:
                                    print(f"[WARNING] Не удалось установить жирность для {part} через XML: {xml_bold_error}")
                            else:
                                # ВАЖНО: Убеждаемся, что значение переменной НЕ жирное (если не PIN-код, issue_date или doc_number)
                                try:
                                    run.font.bold = False
                                except:
                                    pass
                    else:
                        # Обычный текст (не переменная) - это названия переменных, должны быть жирными
                        run = paragraph.add_run(part)
                        try:
                            run.font.name = 'Arial'
                        except:
                            try:
                                run.font.name = font_name
                            except:
                                pass
                        if original_font_size:
                            try:
                                run.font.size = original_font_size
                            except:
                                pass
                        # ВАЖНО: Названия переменных должны быть жирными
                        try:
                            run.font.bold = True
                        except:
                            pass
                        # Дополнительно через XML для гарантии
                        try:
                            from docx.oxml.ns import qn
                            from docx.oxml import OxmlElement
                            if run._element.rPr is None:
                                run._element.get_or_add_rPr()
                            # Убеждаемся, что жирность установлена
                            b_elements = run._element.rPr.findall(qn('w:b'))
                            if not b_elements:
                                b = OxmlElement('w:b')
                                b.set(qn('w:val'), 'true')
                                run._element.rPr.append(b)
                            else:
                                for b in b_elements:
                                    b.set(qn('w:val'), 'true')
                        except Exception as xml_bold_error:
                            print(f"[WARNING] Не удалось установить жирность для названия переменной через XML: {xml_bold_error}")
                        # Сохраняем курсив из оригинального форматирования
                        if original_italic is not None:
                            try:
                                run.font.italic = original_italic
                            except:
                                pass
                
                # Проверяем результат замены
                final_text = paragraph.text
                remaining_placeholders = [key for key in replacements.keys() if key in final_text]
                if remaining_placeholders:
                    print(f"[WARNING] После замены остались плейсхолдеры: {remaining_placeholders}")
                else:
                    print(f"[DEBUG] Все плейсхолдеры успешно заменены. Итоговый текст: {final_text[:100]}")
                    
            except Exception as e:
                print(f"Ошибка при замене плейсхолдеров в параграфе: {e}")
                import traceback
                print(traceback.format_exc())
                try:
                    # Fallback: простая замена без сохранения форматирования
                    simple_text = check_text if 'check_text' in locals() else full_text
                    for key, value in replacements.items():
                        simple_text = simple_text.replace(key, str(value))
                    paragraph.text = simple_text
                    print(f"[INFO] Использован fallback метод замены")
                except Exception as fallback_error:
                    print(f"Ошибка при fallback замене: {fallback_error}")
        
        # Заменяем плейсхолдеры в параграфах
        for paragraph in doc.paragraphs:
            replace_placeholders_with_font(paragraph, replacements)
        
        # Заменяем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_with_font(paragraph, replacements)
        
        # Функция для обработки колонтитула (header или footer)
        def process_header_footer(header_footer, header_footer_name=""):
            """Обрабатывает все параграфы и таблицы в колонтитуле"""
            if not header_footer:
                return
            
            try:
                # Обрабатываем параграфы
                for paragraph in header_footer.paragraphs:
                    try:
                        replace_placeholders_with_font(paragraph, replacements)
                        print(f"[DEBUG] Обработан параграф в {header_footer_name}: {paragraph.text[:50]}")
                    except Exception as e:
                        print(f"[WARNING] Ошибка при обработке параграфа в {header_footer_name}: {e}")
                
                # Обрабатываем таблицы
                for table in header_footer.tables:
                    try:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    try:
                                        replace_placeholders_with_font(paragraph, replacements)
                                        print(f"[DEBUG] Обработан параграф в таблице {header_footer_name}: {paragraph.text[:50]}")
                                    except Exception as e:
                                        print(f"[WARNING] Ошибка при обработке параграфа в таблице {header_footer_name}: {e}")
                    except Exception as e:
                        print(f"[WARNING] Ошибка при обработке таблицы в {header_footer_name}: {e}")
            except Exception as e:
                print(f"[WARNING] Ошибка при обработке {header_footer_name}: {e}")
        
        # Заменяем плейсхолдеры во всех типах колонтитулов
        for section_idx, section in enumerate(doc.sections):
            print(f"[INFO] Обработка секции {section_idx + 1}, колонтитулы:")
            
            # Обычные колонтитулы
            if section.header:
                print(f"[INFO]  - Обычный header найден")
                process_header_footer(section.header, f"секция {section_idx + 1}, header")
            
            if section.footer:
                print(f"[INFO]  - Обычный footer найден")
                process_header_footer(section.footer, f"секция {section_idx + 1}, footer")
            
            # Колонтитулы первой страницы
            try:
                if hasattr(section, 'first_page_header') and section.first_page_header:
                    print(f"[INFO]  - Header первой страницы найден")
                    process_header_footer(section.first_page_header, f"секция {section_idx + 1}, first_page_header")
            except Exception as e:
                print(f"[DEBUG] first_page_header недоступен: {e}")
            
            try:
                if hasattr(section, 'first_page_footer') and section.first_page_footer:
                    print(f"[INFO]  - Footer первой страницы найден")
                    process_header_footer(section.first_page_footer, f"секция {section_idx + 1}, first_page_footer")
            except Exception as e:
                print(f"[DEBUG] first_page_footer недоступен: {e}")
            
            # Колонтитулы четных/нечетных страниц (если включены)
            try:
                if hasattr(section, 'even_page_header') and section.even_page_header:
                    print(f"[INFO]  - Header четных страниц найден")
                    process_header_footer(section.even_page_header, f"секция {section_idx + 1}, even_page_header")
            except Exception as e:
                print(f"[DEBUG] even_page_header недоступен: {e}")
            
            try:
                if hasattr(section, 'even_page_footer') and section.even_page_footer:
                    print(f"[INFO]  - Footer четных страниц найден")
                    process_header_footer(section.even_page_footer, f"секция {section_idx + 1}, even_page_footer")
            except Exception as e:
                print(f"[DEBUG] even_page_footer недоступен: {e}")
            
            try:
                if hasattr(section, 'odd_page_header') and section.odd_page_header:
                    print(f"[INFO]  - Header нечетных страниц найден")
                    process_header_footer(section.odd_page_header, f"секция {section_idx + 1}, odd_page_header")
            except Exception as e:
                print(f"[DEBUG] odd_page_header недоступен: {e}")
            
            try:
                if hasattr(section, 'odd_page_footer') and section.odd_page_footer:
                    print(f"[INFO]  - Footer нечетных страниц найден")
                    process_header_footer(section.odd_page_footer, f"секция {section_idx + 1}, odd_page_footer")
            except Exception as e:
                print(f"[DEBUG] odd_page_footer недоступен: {e}")
            
            # Альтернативный способ доступа к колонтитулам через XML (если стандартный не работает)
            try:
                from docx.oxml.ns import qn
                header_part = section._sectPr.get_or_add_headerReference()
                footer_part = section._sectPr.get_or_add_footerReference()
                
                # Пробуем получить все типы колонтитулов через XML
                header_refs = section._sectPr.findall(qn('w:headerReference'))
                footer_refs = section._sectPr.findall(qn('w:footerReference'))
                
                if header_refs:
                    print(f"[INFO]  - Найдено {len(header_refs)} header references через XML")
                if footer_refs:
                    print(f"[INFO]  - Найдено {len(footer_refs)} footer references через XML")
            except Exception as e:
                print(f"[DEBUG] XML доступ к колонтитулам: {e}")
        
        # Добавляем QR-код
        add_qr_code_to_docx(doc, document_data.get('pin_code', ''), app, document_data.get('uuid', ''))
        
        # Сохраняем заполненный документ
        upload_folder = app.config.get('UPLOAD_FOLDER', 'static/generated_documents') if app else 'static/generated_documents'
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder, exist_ok=True)
        
        document_uuid = document_data.get('uuid', '')
        if not document_uuid:
            document_uuid = str(uuid.uuid4())
        output_path = os.path.join(upload_folder, f"{document_uuid}.docx")
        
        print(f"Сохраняем документ: {output_path}")
        try:
            doc.save(output_path)
            print(f"Документ успешно сохранен: {output_path}")
            
            if os.path.exists(output_path):
                with open(output_path, 'rb') as f:
                    docx_data = f.read()
                
                docx_filename = f"{document_uuid}.docx"
                stored_path = storage_manager.save_file(docx_data, docx_filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                
                if storage_manager.use_minio and os.path.exists(output_path):
                    try:
                        os.remove(output_path)
                    except:
                        pass
                
                return stored_path
            else:
                print(f"ОШИБКА: Файл не был создан: {output_path}")
                return None
        except Exception as save_error:
            print(f"ОШИБКА при сохранении документа: {save_error}")
            import traceback
            print(traceback.format_exc())
            return None
        
    except Exception as e:
        import traceback
        error_msg = f"Ошибка при заполнении DOCX шаблона: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return None


def create_default_docx_template():
    """Создает стандартный DOCX шаблон если его нет"""
    doc = Document()
    
    # Настраиваем размеры страницы A4 и минимальные отступы
    from docx.shared import Cm
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)  # A4 высота
        section.page_width = Cm(21.0)   # A4 ширина
        section.left_margin = Cm(1.0)   # Минимальные отступы
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
    
    try:
        title = doc.add_heading("Ta'lim olayotgan shaxslar uchun mehnatga layoqatsizlik ma'lumotnomasi", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        title = doc.add_paragraph("Ta'lim olayotgan shaxslar uchun mehnatga layoqatsizlik ma'lumotnomasi")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].font.bold = True
            title.runs[0].font.name = DOCX_FONT_NAME
    
    org_para = doc.add_paragraph("O'zbekiston Respublikasi Sog'liqni saqlash vazirligi")
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_para = doc.add_paragraph("4 - sonli oilaviy poliklinika")
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Ro'yhatga olingan sana: {{issue_date}}")
    doc.add_paragraph(f"No {{doc_number}}")
    
    try:
        doc.add_heading('Vaqtincha mehnatga layoqatsiz fuqaro haqidagi ma\'lumotlar:', 1)
    except:
        para = doc.add_paragraph('Vaqtincha mehnatga layoqatsiz fuqaro haqidagi ma\'lumotlar:')
        para.runs[0].font.bold = True
        para.runs[0].font.name = DOCX_FONT_NAME
    
    doc.add_paragraph(f"FISH: {{patient_name}}")
    doc.add_paragraph(f"Jinsi: {{gender}}")
    doc.add_paragraph(f"JShShIR: {{jshshir}}")
    doc.add_paragraph(f"Yoshi: {{age}}")
    doc.add_paragraph(f"Yashash manzili: {{address}}")
    doc.add_paragraph(f"Biriktirilgan tibbiy muassasa: {{attached_medical_institution}}")
    doc.add_paragraph(f"Tashxis (KXT-10 kodi va Nomi): {{diagnosis_icd10_code}}: {{diagnosis}}")
    doc.add_paragraph(f"Yakuniy tashxis (Nomi va KXT-10 kodi): {{final_diagnosis_icd10_code}}: {{final_diagnosis}}")
    doc.add_paragraph(f"Ro'yhatga olingan sana: {{issue_date}}")
    doc.add_paragraph(f"No: {{doc_number}}")
    doc.add_paragraph(f"Davolovchi shifokor FISH: {{doctor_name}}")
    doc.add_paragraph(f"Bo'lim boshlig'i (mas'ul shaxs) FISH: {{department_head_name}}")
    doc.add_paragraph(f"Ishdan ozod etilgan kunlar: {{days_off_period}}")
    
    try:
        doc.add_heading('Shifokor ma\'lumotlari:', 1)
    except:
        para = doc.add_paragraph('Shifokor ma\'lumotlari:')
        if para.runs:
            para.runs[0].font.bold = True
            para.runs[0].font.name = DOCX_FONT_NAME
        else:
            run = para.add_run('Shifokor ma\'lumotlari:')
            run.font.bold = True
            run.font.name = DOCX_FONT_NAME
    
    doc.add_paragraph(f"FISH: {{doctor_name}}")
    doc.add_paragraph(f"Lavozim: {{doctor_position}}")
    doc.add_paragraph(f"PIN-kod: {{pin_code}}")
    
    return doc


def add_qr_code_to_docx(doc, pin_code, app=None, document_uuid=None):
    """Добавляет QR-код в DOCX документ с PIN-кодом слева"""
    print(f"[QR_PIN_LAYOUT] ===== Начало добавления QR-кода и PIN-кода =====")
    print(f"[QR_PIN_LAYOUT] PIN-код: {pin_code}, UUID: {document_uuid}")
    try:
        # Генерируем URL для QR-кода: используем FRONTEND_URL из конфига
        from config import FRONTEND_URL
        
        if document_uuid:
            # Убираем trailing slash если есть
            frontend_url = FRONTEND_URL.rstrip('/')
            qr_url = f"{frontend_url}/access/{document_uuid}"
        else:
            # Fallback на старый способ, если UUID не передан
            try:
                from flask import url_for
                qr_url = url_for('verify_document', _external=True)
            except RuntimeError:
                qr_url = '/verify'
        
        print(f"[QR_PIN_LAYOUT] URL для QR-кода: {qr_url}")
        qr_img = generate_qr_code(qr_url)
        
        if qr_img is None:
            raise Exception("generate_qr_code() вернул None")
        
        if not isinstance(qr_img, Image.Image):
            raise Exception(f"generate_qr_code() вернул не PIL.Image, а {type(qr_img)}")
        
        print(f"[QR_PIN_LAYOUT] QR-код сгенерирован, размер: {qr_img.size if qr_img else 'None'}, режим: {qr_img.mode if qr_img else 'None'}")
        
        upload_folder = app.config.get('UPLOAD_FOLDER', 'static/generated_documents') if app else 'static/generated_documents'
        qr_temp_path = os.path.join(upload_folder, f'qr_temp_{pin_code}.png')
        if not os.path.exists(os.path.dirname(qr_temp_path)):
            os.makedirs(os.path.dirname(qr_temp_path), exist_ok=True)
        
        # Убеждаемся, что изображение в правильном формате перед сохранением
        if qr_img.mode != 'RGBA':
            try:
                qr_img = qr_img.convert('RGBA')
                print(f"[QR_PIN_LAYOUT] QR-код конвертирован в RGBA перед сохранением")
            except Exception as conv_error:
                print(f"[WARNING] Не удалось конвертировать QR-код в RGBA: {conv_error}")
        
        # Сохраняем изображение
        try:
            with open(qr_temp_path, 'wb') as f:
                qr_img.save(f, 'PNG', optimize=True)
            print(f"[QR_PIN_LAYOUT] QR-код сохранен во временный файл: {qr_temp_path}")
            
            # Проверяем, что файл действительно создан и имеет размер
            if os.path.exists(qr_temp_path):
                file_size = os.path.getsize(qr_temp_path)
                print(f"[QR_PIN_LAYOUT] Файл QR-кода создан, размер: {file_size} байт")
                if file_size == 0:
                    raise Exception(f"Файл QR-кода пустой: {qr_temp_path}")
            else:
                raise Exception(f"Файл QR-кода не был создан: {qr_temp_path}")
        except Exception as save_error:
            print(f"[ERROR] Ошибка при сохранении QR-кода: {save_error}")
            import traceback
            print(traceback.format_exc())
            raise
        
        qr_placeholder_found = False
        
        def create_pin_qr_table():
            """Создает таблицу с PIN-кодом слева и QR-кодом справа"""
            print(f"[QR_PIN_LAYOUT] Создаем таблицу для PIN ({pin_code}) и QR-кода")
            # Создаем таблицу с двумя колонками: PIN-код слева, QR-код справа
            table = doc.add_table(rows=1, cols=2)
            # Убираем стиль таблицы чтобы не было видимых границ
            table.style = None
            
            # Настраиваем ширину колонок - делаем компактнее
            from docx.shared import Cm, Pt
            # Для 4-значного PIN-кода с размером шрифта 20pt нужно минимум 1.8-2.0 см
            table.columns[0].width = Cm(2.0)  # Увеличена ширина для PIN-кода (чтобы точно не переносился)
            # Увеличиваем ширину для QR-кода, чтобы он не обрезался (1.2 дюйма = ~3.05 см, нужно минимум 3.5 см с отступами)
            table.columns[1].width = Cm(3.8)  # Увеличена ширина для QR-кода (чтобы не обрезался)
            print(f"[QR_PIN_LAYOUT] Ширина колонок: PIN={Cm(2.0)}, QR={Cm(3.8)}")
            
            # Убираем отступы в ячейках для компактности
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # Ячейка 1: PIN-код слева (большой жирный текст)
            cell_pin = table.rows[0].cells[0]
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            cell_pin.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # Убираем отступы в ячейке через XML
            tc_pr = cell_pin._element.tcPr
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                cell_pin._element.append(tc_pr)
            # Отключаем перенос текста (noWrap) чтобы PIN-код не разбивался на строки
            no_wrap = OxmlElement('w:noWrap')
            tc_pr.append(no_wrap)
            # Устанавливаем минимальные отступы (используем положительные значения для совместимости)
            tc_mar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom']:
                margin_elem = OxmlElement(f'w:{margin}')
                margin_elem.set(qn('w:w'), '0')
                margin_elem.set(qn('w:type'), 'dxa')
                tc_mar.append(margin_elem)
            # Отступ справа - минимальный для сближения с QR-кодом (в пикселях: ~2-3px = ~15-20 twips)
            right_margin = OxmlElement('w:right')
            right_margin.set(qn('w:w'), '20')  # Минимальный отступ ~2-3px для сближения
            right_margin.set(qn('w:type'), 'dxa')
            tc_mar.append(right_margin)
            print(f"[QR_PIN_LAYOUT] Отступы ячейки PIN установлены (все нулевые для совместимости)")
            # Удаляем старый tcMar если есть
            old_tc_mar = tc_pr.find(qn('w:tcMar'))
            if old_tc_mar is not None:
                tc_pr.remove(old_tc_mar)
            tc_pr.append(tc_mar)
            para_pin = cell_pin.paragraphs[0]
            para_pin.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравниваем по левому краю ячейки
            # Убираем отступы параграфа
            para_pin_format = para_pin.paragraph_format
            para_pin_format.space_before = Pt(0)
            para_pin_format.space_after = Pt(0)
            para_pin_format.left_indent = Pt(0)  # Нет отступа слева - прижимаем к левому краю
            para_pin_format.right_indent = Pt(0)
            # Предотвращаем перенос текста на уровне параграфа
            para_pin_format.widow_control = False
            para_pin_format.keep_together = True
            # Очищаем параграф перед добавлением (на случай если там что-то осталось)
            para_pin.clear()
            # Добавляем PIN-код как один run - убеждаемся что он в одной строке
            pin_text = str(pin_code).strip()  # Убираем пробелы
            # Убеждаемся, что PIN-код в одной строке - заменяем все пробелы и переносы
            pin_text = pin_text.replace('\n', '').replace('\r', '').replace(' ', '')
            run_pin = para_pin.add_run(pin_text)
            run_pin.font.size = Pt(20)  # Размер шрифта для PIN-кода
            run_pin.font.bold = True
            # Добавляем атрибут noBreak для предотвращения переноса внутри run
            r_pr = run_pin._element.get_or_add_rPr()
            no_break = OxmlElement('w:noBreak')
            r_pr.append(no_break)
            print(f"[QR_PIN_LAYOUT] PIN-код добавлен как одна строка: '{pin_text}'")
            try:
                run_pin.font.name = 'Arial'
            except:
                try:
                    run_pin.font.name = 'DejaVu Sans'
                except:
                    pass
            print(f"[QR_PIN_LAYOUT] PIN-код добавлен в ячейку 0 (слева): {pin_code}")
            
            # Ячейка 2: QR-код справа
            print(f"[QR_PIN_LAYOUT] Добавляем QR-код в ячейку 1 (справа)")
            cell_qr = table.rows[0].cells[1]
            cell_qr.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # Убираем отступы в ячейке через XML
            tc_pr_qr = cell_qr._element.tcPr
            if tc_pr_qr is None:
                tc_pr_qr = OxmlElement('w:tcPr')
                cell_qr._element.append(tc_pr_qr)
            # Устанавливаем минимальные отступы, чтобы QR-код не обрезался
            # Используем положительные значения для совместимости между системами
            tc_mar_qr = OxmlElement('w:tcMar')
            for margin in ['top', 'bottom', 'right']:
                margin_elem = OxmlElement(f'w:{margin}')
                margin_elem.set(qn('w:w'), '0')
                margin_elem.set(qn('w:type'), 'dxa')
                tc_mar_qr.append(margin_elem)
            # Отступ слева - минимальный для сближения с PIN-кодом (в пикселях: ~2-3px = ~15-20 twips)
            left_margin = OxmlElement('w:left')
            left_margin.set(qn('w:w'), '20')  # Минимальный отступ ~2-3px для сближения с PIN
            left_margin.set(qn('w:type'), 'dxa')
            tc_mar_qr.append(left_margin)
            print(f"[QR_PIN_LAYOUT] Отступы ячейки QR установлены (все нулевые для совместимости)")
            # Удаляем старый tcMar если есть
            old_tc_mar_qr = tc_pr_qr.find(qn('w:tcMar'))
            if old_tc_mar_qr is not None:
                tc_pr_qr.remove(old_tc_mar_qr)
            tc_pr_qr.append(tc_mar_qr)
            para_qr = cell_qr.paragraphs[0]
            # Выравниваем по левому краю ячейки
            para_qr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Убираем отступы параграфа
            para_qr_format = para_qr.paragraph_format
            para_qr_format.space_before = Pt(0)
            para_qr_format.space_after = Pt(0)
            para_qr_format.left_indent = Pt(0)  # Нет отступа слева - прижимаем к левому краю ячейки
            para_qr_format.right_indent = Pt(0)
            run_qr = para_qr.add_run()
            # Размер QR-кода - оптимизирован для ячейки шириной 3.8 см
            # 1.1 дюйма = ~2.79 см, что помещается в ячейку 3.8 см с отступами
            qr_width_inches = 1.1  # Оптимизирован для ячейки 3.8 см (чтобы не обрезался)
            run_qr.add_picture(qr_temp_path, width=Inches(qr_width_inches))
            print(f"[QR_PIN_LAYOUT] QR-код добавлен в ячейку 1, размер: {qr_width_inches} дюймов (~{qr_width_inches * 2.54:.2f} см)")
            
            # Выравниваем таблицу по левому краю
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(f"[QR_PIN_LAYOUT] Таблица создана: PIN слева (ячейка 0), QR справа (ячейка 1)")
            
            return table
        
        def replace_qr_placeholder(paragraph, in_table_cell=False, cell=None):
            """Заменяет плейсхолдер QR-кода на таблицу с PIN-кодом слева и QR-кодом справа"""
            nonlocal qr_placeholder_found
            if '{{qr_code}}' in paragraph.text or '{{pin_code_with_qr}}' in paragraph.text:
                qr_placeholder_found = True
                
                if in_table_cell and cell:
                    # Если мы в ячейке таблицы, создаем вложенную таблицу
                    print(f"[QR_PIN_LAYOUT] Найден плейсхолдер в ячейке таблицы, создаем вложенную таблицу")
                    # Полностью очищаем ячейку - удаляем все параграфы
                    # Это важно, так как PIN-код мог быть уже заменен и разбит на несколько параграфов
                    for paragraph in cell.paragraphs[:]:
                        p_element = paragraph._element
                        p_element.getparent().remove(p_element)
                    print(f"[QR_PIN_LAYOUT] Ячейка полностью очищена от всех параграфов")
                    
                    # Создаем вложенную таблицу в ячейке
                    inner_table = cell.add_table(rows=1, cols=2)
                    # Убираем стиль таблицы
                    inner_table.style = None
                    # Выравниваем таблицу по левому краю родительской ячейки
                    inner_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    from docx.shared import Cm, Pt
                    # Уменьшаем ширину колонок для компактности - делаем их ближе друг к другу
                    # Увеличиваем ширину для PIN-кода, чтобы он не переносился на две строки
                    # Для 4-значного PIN-кода с размером шрифта 20pt нужно минимум 1.8-2.0 см
                    inner_table.columns[0].width = Cm(2.0)  # Увеличена ширина для PIN-кода (чтобы точно не переносился)
                    # Увеличиваем ширину для QR-кода, чтобы он не обрезался (1.2 дюйма = ~3.05 см, нужно минимум 3.5 см с отступами)
                    inner_table.columns[1].width = Cm(3.8)  # Увеличена ширина для QR-кода (чтобы не обрезался)
                    print(f"[QR_PIN_LAYOUT] Вложенная таблица создана: PIN={Cm(2.0)}, QR={Cm(3.8)}, выравнивание: LEFT")
                    
                    # Убираем отступы в ячейках
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    
                    # Ячейка 1: PIN-код
                    inner_cell_pin = inner_table.rows[0].cells[0]
                    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
                    inner_cell_pin.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    # Убираем отступы через XML
                    tc_pr_pin = inner_cell_pin._element.tcPr
                    if tc_pr_pin is None:
                        tc_pr_pin = OxmlElement('w:tcPr')
                        inner_cell_pin._element.append(tc_pr_pin)
                    # Отключаем перенос текста (noWrap) чтобы PIN-код не разбивался на строки
                    no_wrap_pin = OxmlElement('w:noWrap')
                    tc_pr_pin.append(no_wrap_pin)
                    # Устанавливаем минимальные отступы (используем положительные значения для совместимости)
                    tc_mar_pin = OxmlElement('w:tcMar')
                    for margin in ['top', 'left', 'bottom']:
                        margin_elem = OxmlElement(f'w:{margin}')
                        margin_elem.set(qn('w:w'), '0')
                        margin_elem.set(qn('w:type'), 'dxa')
                        tc_mar_pin.append(margin_elem)
                    # Отступ справа - минимальный для сближения с QR-кодом (в пикселях: ~2-3px = ~15-20 twips)
                    right_margin_pin = OxmlElement('w:right')
                    right_margin_pin.set(qn('w:w'), '20')  # Минимальный отступ ~2-3px для сближения
                    right_margin_pin.set(qn('w:type'), 'dxa')
                    tc_mar_pin.append(right_margin_pin)
                    old_tc_mar_pin = tc_pr_pin.find(qn('w:tcMar'))
                    if old_tc_mar_pin is not None:
                        tc_pr_pin.remove(old_tc_mar_pin)
                    tc_pr_pin.append(tc_mar_pin)
                    para_pin = inner_cell_pin.paragraphs[0]
                    para_pin.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравниваем по левому краю ячейки
                    # Убираем отступы параграфа
                    para_pin_format = para_pin.paragraph_format
                    para_pin_format.space_before = Pt(0)
                    para_pin_format.space_after = Pt(0)
                    para_pin_format.left_indent = Pt(0)  # Нет отступа слева - прижимаем к левому краю
                    para_pin_format.right_indent = Pt(0)
                    # Предотвращаем перенос текста на уровне параграфа
                    para_pin_format.widow_control = False
                    para_pin_format.keep_together = True
                    # Добавляем PIN-код как один run - убеждаемся что он не переносится
                    # Важно: добавляем PIN-код как одну строку без пробелов и переносов
                    pin_text = str(pin_code).strip()  # Убираем пробелы
                    # Убеждаемся, что PIN-код в одной строке - заменяем все пробелы и переносы
                    pin_text = pin_text.replace('\n', '').replace('\r', '').replace(' ', '')
                    
                    # Очищаем параграф перед добавлением (на случай если там что-то осталось)
                    para_pin.clear()
                    
                    # Добавляем PIN-код как один run
                    run_pin = para_pin.add_run(pin_text)
                    run_pin.font.size = Pt(20)  # Размер шрифта для PIN-кода
                    run_pin.font.bold = True
                    # Добавляем атрибут noBreak для предотвращения переноса внутри run
                    r_pr = run_pin._element.get_or_add_rPr()
                    no_break = OxmlElement('w:noBreak')
                    r_pr.append(no_break)
                    print(f"[QR_PIN_LAYOUT] PIN-код добавлен как одна строка: '{pin_text}'")
                    from docx.oxml.ns import qn as qn_ns
                    r_pr = run_pin._element.get_or_add_rPr()
                    no_break = OxmlElement('w:noBreak')
                    r_pr.append(no_break)
                    try:
                        run_pin.font.name = 'Arial'
                    except:
                        try:
                            run_pin.font.name = 'DejaVu Sans'
                        except:
                            pass
                    print(f"[QR_PIN_LAYOUT] PIN-код добавлен во вложенную таблицу, ячейка 0: {pin_code}")
                    
                    # Ячейка 2: QR-код
                    inner_cell_qr = inner_table.rows[0].cells[1]
                    inner_cell_qr.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    # Убираем отступы через XML
                    tc_pr_qr = inner_cell_qr._element.tcPr
                    if tc_pr_qr is None:
                        tc_pr_qr = OxmlElement('w:tcPr')
                        inner_cell_qr._element.append(tc_pr_qr)
                    # Устанавливаем минимальные отступы, чтобы QR-код не обрезался (используем положительные значения)
                    tc_mar_qr = OxmlElement('w:tcMar')
                    for margin in ['top', 'bottom', 'right']:
                        margin_elem = OxmlElement(f'w:{margin}')
                        margin_elem.set(qn('w:w'), '0')
                        margin_elem.set(qn('w:type'), 'dxa')
                        tc_mar_qr.append(margin_elem)
                    # Отступ слева - минимальный для сближения с PIN-кодом (в пикселях: ~2-3px = ~15-20 twips)
                    left_margin = OxmlElement('w:left')
                    left_margin.set(qn('w:w'), '20')  # Минимальный отступ ~2-3px для сближения с PIN
                    left_margin.set(qn('w:type'), 'dxa')
                    tc_mar_qr.append(left_margin)
                    old_tc_mar_qr = tc_pr_qr.find(qn('w:tcMar'))
                    if old_tc_mar_qr is not None:
                        tc_pr_qr.remove(old_tc_mar_qr)
                    tc_pr_qr.append(tc_mar_qr)
                    para_qr = inner_cell_qr.paragraphs[0]
                    # Выравниваем по левому краю ячейки
                    para_qr.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Убираем отступы параграфа
                    para_qr_format = para_qr.paragraph_format
                    para_qr_format.space_before = Pt(0)
                    para_qr_format.space_after = Pt(0)
                    para_qr_format.left_indent = Pt(0)  # Нет отступа слева - прижимаем к левому краю ячейки
                    para_qr_format.right_indent = Pt(0)
                    run_qr = para_qr.add_run()
                    # Размер QR-кода - оптимизирован для ячейки шириной 3.8 см
                    # 1.1 дюйма = ~2.79 см, что помещается в ячейку 3.8 см с отступами
                    qr_width_inches = 1.1  # Оптимизирован для ячейки 3.8 см (чтобы не обрезался)
                    run_qr.add_picture(qr_temp_path, width=Inches(qr_width_inches))
                    print(f"[QR_PIN_LAYOUT] QR-код добавлен во вложенную таблицу, ячейка 1, размер: {qr_width_inches} дюймов (~{qr_width_inches * 2.54:.2f} см)")
                    print(f"[QR_PIN_LAYOUT] Вложенная таблица завершена: PIN слева (ячейка 0), QR справа (ячейка 1)")
                else:
                    # Если мы в обычном параграфе, создаем таблицу
                    # Получаем родительский элемент параграфа
                    parent = paragraph._element.getparent()
                    
                    # Создаем таблицу
                    table = create_pin_qr_table()
                    
                    # Вставляем таблицу перед параграфом
                    table_element = table._element
                    parent.insert(parent.index(paragraph._element), table_element)
                    
                    # Удаляем оригинальный параграф
                    parent.remove(paragraph._element)
                
                return True
            return False
        
        for paragraph in doc.paragraphs:
            if replace_qr_placeholder(paragraph, in_table_cell=False):
                break
        
        if not qr_placeholder_found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if replace_qr_placeholder(paragraph, in_table_cell=True, cell=cell):
                                qr_placeholder_found = True
                                break
                        if qr_placeholder_found:
                            break
                    if qr_placeholder_found:
                        break
                if qr_placeholder_found:
                    break
        
        if not qr_placeholder_found:
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if replace_qr_placeholder(paragraph, in_table_cell=False):
                            qr_placeholder_found = True
                            break
                    if not qr_placeholder_found:
                        for table in section.header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if replace_qr_placeholder(paragraph, in_table_cell=True, cell=cell):
                                            qr_placeholder_found = True
                                            break
                                    if qr_placeholder_found:
                                        break
                                if qr_placeholder_found:
                                    break
                            if qr_placeholder_found:
                                break
                
                if qr_placeholder_found:
                    break
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if replace_qr_placeholder(paragraph, in_table_cell=False):
                            qr_placeholder_found = True
                            break
                    if not qr_placeholder_found:
                        for table in section.footer.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if replace_qr_placeholder(paragraph, in_table_cell=True, cell=cell):
                                            qr_placeholder_found = True
                                            break
                                    if qr_placeholder_found:
                                        break
                                if qr_placeholder_found:
                                    break
                            if qr_placeholder_found:
                                break
                
                if qr_placeholder_found:
                    break
        
        if not qr_placeholder_found:
            # Если плейсхолдер не найден, добавляем таблицу с PIN-кодом и QR-кодом в конец документа
            print(f"[QR_PIN_LAYOUT] Плейсхолдер не найден, добавляем таблицу в конец документа")
            create_pin_qr_table()
        else:
            print(f"[QR_PIN_LAYOUT] Плейсхолдер найден и заменен")
        
        if os.path.exists(qr_temp_path):
            os.remove(qr_temp_path)
            print(f"[QR_PIN_LAYOUT] Временный файл QR-кода удален")
        
        print(f"[QR_PIN_LAYOUT] ===== Завершение добавления QR-кода и PIN-кода =====")
            
    except Exception as e:
        print(f"[ERROR] Ошибка при добавлении QR-кода: {e}")
        import traceback
        print(traceback.format_exc())

