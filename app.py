# Настройка кодировки для Windows консоли
import sys
import io
if sys.platform == 'win32':
    # Устанавливаем UTF-8 для stdout/stderr в Windows
    if sys.stdout.encoding != 'utf-8':
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    if sys.stderr.encoding != 'utf-8':
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from flask import Flask, render_template, request, jsonify, send_file, url_for
from datetime import datetime
import qrcode
from io import BytesIO
import os
import random
import string
import re
import uuid
import base64
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
import psycopg2
from psycopg2.extras import RealDictCursor
from psycopg2 import pool
from config import DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASSWORD, DB_SSLMODE, SECRET_KEY, UPLOAD_FOLDER, DOC_NUMBER_PREFIX, DOC_NUMBER_FORMAT, DOCX_FONT_NAME
from storage import storage_manager
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Попытка импортировать docx2pdf (может не работать без LibreOffice/Word)
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("Предупреждение: docx2pdf не доступен. Конвертация DOCX->PDF будет использовать альтернативный метод.")

# Попытка импортировать mammoth и weasyprint для конвертации DOCX->PDF без LibreOffice
try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    print("Предупреждение: mammoth не доступен. Установите: pip install mammoth")

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    print(f"Предупреждение: weasyprint не доступен ({type(e).__name__}). Будет использован альтернативный метод конвертации.")

app = Flask(__name__)
app.config['SECRET_KEY'] = SECRET_KEY
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMPLATE_FOLDER'] = 'templates/docx_templates'

# Инициализация пула подключений к PostgreSQL
def get_db_connection():
    """Создает подключение к PostgreSQL"""
    if not DB_PASSWORD:
        raise ValueError("DB_PASSWORD не установлен. Проверьте файл .env")
    return psycopg2.connect(
        host=DB_HOST,
        port=DB_PORT,
        database=DB_NAME,
        user=DB_USER,
        password=DB_PASSWORD,
        sslmode=DB_SSLMODE
    )

# Создаем пул подключений (отложенная инициализация)
db_pool = None

def init_db_pool():
    """Инициализирует пул подключений к БД"""
    global db_pool
    if db_pool is None:
        try:
            from psycopg2 import pool as psycopg2_pool
            if not DB_PASSWORD:
                print("WARNING: DB_PASSWORD не установлен. Проверьте файл .env")
                return None
            
            db_pool = psycopg2_pool.SimpleConnectionPool(
                minconn=1,
                maxconn=10,
                host=DB_HOST,
                port=DB_PORT,
                database=DB_NAME,
                user=DB_USER,
                password=DB_PASSWORD,
                sslmode=DB_SSLMODE
            )
            print("OK: Подключение к PostgreSQL установлено")
            return db_pool
        except Exception as e:
            print(f"ERROR: Ошибка подключения к PostgreSQL: {e}")
            db_pool = None
            return None
    return db_pool

# Создаем директории если их нет
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['TEMPLATE_FOLDER'], exist_ok=True)


# Функции для работы с БД
def db_query(query, params=None, fetch_one=False, fetch_all=False):
    """Выполняет SQL запрос к БД"""
    conn = None
    try:
        # Инициализируем пул при первом использовании
        pool = init_db_pool()
        if pool:
            conn = pool.getconn()
        else:
            conn = get_db_connection()
        
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        cursor.execute(query, params)
        
        if fetch_one:
            result = cursor.fetchone()
            conn.commit()
        elif fetch_all:
            result = cursor.fetchall()
            conn.commit()
        else:
            conn.commit()
            result = cursor.rowcount
        
        cursor.close()
        return result
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Ошибка БД: {e}")
        raise e
    finally:
        if conn:
            pool = init_db_pool()
            if pool:
                pool.putconn(conn)
            else:
                conn.close()


def db_insert(table, data):
    """Вставляет запись в таблицу и возвращает созданную запись"""
    columns = ', '.join(data.keys())
    placeholders = ', '.join(['%s'] * len(data))
    values = list(data.values())
    
    query = f"""
        INSERT INTO {table} ({columns})
        VALUES ({placeholders})
        RETURNING *
    """
    
    print(f"DEBUG db_insert: query = {query[:200]}...")
    print(f"DEBUG db_insert: values count = {len(values)}")
    
    try:
        result = db_query(query, values, fetch_one=True)
        print(f"DEBUG db_insert: result = {result}")
        if result:
            result_dict = dict(result)
            print(f"DEBUG db_insert: result_dict keys = {list(result_dict.keys())}")
            return result_dict
        else:
            print("WARNING: db_insert вернул None")
            return None
    except Exception as e:
        print(f"ERROR db_insert: {e}")
        import traceback
        print(traceback.format_exc())
        raise


def db_select(table, where_clause=None, params=None, fetch_one=False):
    """Выбирает записи из таблицы"""
    query = f"SELECT * FROM {table}"
    if where_clause:
        query += f" WHERE {where_clause}"
    
    if fetch_one:
        result = db_query(query, params, fetch_one=True)
        if result:
            return dict(result)
        return None
    else:
        results = db_query(query, params, fetch_all=True)
        if results:
            return [dict(row) for row in results]
        return []


def db_update(table, data, where_clause, where_params):
    """Обновляет запись в таблице"""
    set_clause = ', '.join([f"{key} = %s" for key in data.keys()])
    values = list(data.values()) + list(where_params)
    
    query = f"""
        UPDATE {table}
        SET {set_clause}
        WHERE {where_clause}
    """
    
    db_query(query, values)


def generate_document_number():
    """Генерирует уникальный номер документа в формате: № 01ВШ XXXXXXX"""
    max_attempts = 100
    
    # Префикс организации из конфига
    prefix = DOC_NUMBER_PREFIX
    
    for _ in range(max_attempts):
        # Генерируем 9-значный номер в зависимости от формата
        if DOC_NUMBER_FORMAT == 'date':
            # Формат: DDMMYY + 3 цифры порядкового номера (рекомендуется)
            now = datetime.now()
            date_part = f"{now.day:02d}{now.month:02d}{now.year % 100:02d}"
            serial_part = random.randint(100, 999)
            number_part = int(f"{date_part}{serial_part}")
        elif DOC_NUMBER_FORMAT == 'random':
            # Случайный 9-значный номер
            number_part = random.randint(100000000, 999999999)
        else:
            # Sequential или по умолчанию - timestamp
            number_part = int(datetime.now().timestamp()) % 1000000000
        
        # Формируем полный номер: № 01ВШ XXXXXXX
        doc_number = f"№ {prefix} {number_part:09d}"
        
        # Проверяем уникальность в БД
        result = db_select('documents', 'doc_number = %s', [doc_number], fetch_one=True)
        if not result:
            return doc_number
    
    # Если не удалось найти уникальный номер, используем timestamp
    timestamp = int(datetime.now().timestamp())
    return f"№ {prefix} {timestamp % 1000000000:09d}"


def generate_pin_code():
    """Генерирует 4-значный PIN-код"""
    return ''.join(random.choices(string.digits, k=4))


def generate_qr_code(url):
    """Генерирует QR-код для заданного URL"""
    # Используем числовое значение вместо константы для совместимости
    ERROR_CORRECT_L = 1  # Уровень коррекции ошибок L (Low)
    qr = qrcode.QRCode(
        version=1,
        error_correction=ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    return img


def generate_pdf_document(document_data):
    """Генерирует PDF документ на основе данных из БД"""
    if not document_data:
        return None
    
    # Создаем директорию если её нет
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    
    # Путь к файлу - используем UUID
    document_uuid = document_data.get('uuid', '')
    if not document_uuid:
        # Если UUID нет, генерируем новый
        document_uuid = str(uuid.uuid4())
    filename = f"{document_uuid}.pdf"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    # Создаем PDF
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    
    # Заголовок
    c.setFont("Helvetica-Bold", 10)
    c.drawString(50, height - 50, "O'zbekiston Respublikasi Sog'liqni")
    c.drawString(50, height - 65, "saqlash vazirlig")
    c.drawString(50, height - 80, "4 - sonli oilaviy poliklinika")
    
    # Название документа
    c.setFont("Helvetica-Bold", 12)
    title = "Ta'lim olayotgan shaxslar uchun mehnatga layoqatsizlik ma'lumotnomasi"
    c.drawCentredString(width/2, height - 120, title)
    
    # Дата и номер
    c.setFont("Helvetica", 10)
    issue_date = document_data.get('issue_date')
    if issue_date:
        if isinstance(issue_date, str):
            try:
                dt = datetime.fromisoformat(issue_date.replace('Z', '+00:00'))
                issue_date_str = dt.strftime("%d.%m.%Y %H:%M")
            except:
                try:
                    dt = datetime.strptime(issue_date[:16], "%Y-%m-%dT%H:%M")
                    issue_date_str = dt.strftime("%d.%m.%Y %H:%M")
                except:
                    issue_date_str = issue_date[:10] if len(issue_date) >= 10 else issue_date
        else:
            issue_date_str = issue_date.strftime("%d.%m.%Y %H:%M")
    else:
        issue_date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    
    c.drawCentredString(width/2, height - 140, f"Ro'yhatga olingan sana: {issue_date_str}")
    c.drawCentredString(width/2, height - 155, f"No {doc_number}")
    
    # Информация о пациенте
    y_position = height - 200
    c.setFont("Helvetica", 10)
    
    patient_name = document_data.get('patient_name')
    if patient_name:
        c.drawString(50, y_position, f"FISH: {patient_name}")
        y_position -= 20
    
    gender = document_data.get('gender')
    if gender:
        c.drawString(50, y_position, f"Jinsi: {gender}")
        y_position -= 20
    
    jshshir = document_data.get('jshshir')
    if jshshir:
        c.drawString(50, y_position, f"JShShIR: {jshshir}")
        y_position -= 20
        
    age = document_data.get('age')
    if age:
        c.drawString(50, y_position, f"Yoshi: {age}")
        y_position -= 20
    
    address = document_data.get('address')
    if address:
        # Разбиваем длинный адрес на несколько строк если нужно
        address_lines = address[:80] if len(address) <= 80 else address[:80] + "..."
        c.drawString(50, y_position, f"Yashash manzili: {address_lines}")
        y_position -= 20
    
    attached_medical = document_data.get('attached_medical_institution')
    if attached_medical:
        c.drawString(50, y_position, f"Biriktirilgan tibbiy muassasa: {attached_medical}")
        y_position -= 20
    
    diagnosis = document_data.get('diagnosis')
    diagnosis_code = document_data.get('diagnosis_icd10_code')
    if diagnosis:
        diagnosis_text = f"Tashxis (KXT-10 kodi va Nomi): {diagnosis_code}: {diagnosis}" if diagnosis_code else f"Tashxis: {diagnosis}"
        c.drawString(50, y_position, diagnosis_text)
        y_position -= 20
    
    final_diagnosis = document_data.get('final_diagnosis')
    final_diagnosis_code = document_data.get('final_diagnosis_icd10_code')
    if final_diagnosis:
        final_text = f"Yakuniy tashxis (Nomi va KXT-10 kodi): {final_diagnosis_code}: {final_diagnosis}" if final_diagnosis_code else f"Yakuniy tashxis: {final_diagnosis}"
        c.drawString(50, y_position, final_text)
        y_position -= 20
    
    doctor_name = document_data.get('doctor_name')
    if doctor_name:
        c.drawString(50, y_position, f"Davolovchi shifokor FISH: {doctor_name}")
        y_position -= 20
    
    department_head = document_data.get('department_head_name')
    if department_head:
        c.drawString(50, y_position, f"Bo'lim boshlig'i (mas'ul shaxs) FISH: {department_head}")
        y_position -= 20
    
    days_off_from = document_data.get('days_off_from')
    days_off_to = document_data.get('days_off_to')
    if days_off_from or days_off_to:
        if days_off_from and days_off_to:
            try:
                from_str = days_off_from[:10] if isinstance(days_off_from, str) else days_off_from.strftime("%d.%m.%Y")
                to_str = days_off_to[:10] if isinstance(days_off_to, str) else days_off_to.strftime("%d.%m.%Y")
                c.drawString(50, y_position, f"Ishdan ozod etilgan kunlar: {from_str} - {to_str}")
            except:
                c.drawString(50, y_position, f"Ishdan ozod etilgan kunlar: {days_off_from} - {days_off_to}")
        elif days_off_from:
            try:
                from_str = days_off_from[:10] if isinstance(days_off_from, str) else days_off_from.strftime("%d.%m.%Y")
                c.drawString(50, y_position, f"Ishdan ozod etilgan kunlar: {from_str}")
            except:
                c.drawString(50, y_position, f"Ishdan ozod etilgan kunlar: {days_off_from}")
        y_position -= 20
    
    
    # Генерируем QR-код с UUID для прямого доступа к документу
    document_uuid = document_data.get('uuid', '')
    if document_uuid:
        # QR-код ведет на страницу верификации с UUID
        qr_url = url_for('verify_by_uuid', uuid=document_uuid, _external=True)
    else:
        qr_url = url_for('verify_document', _external=True)
    qr_img = generate_qr_code(qr_url)
    
    # Сохраняем QR-код во временный файл
    qr_buffer = BytesIO()
    qr_img.save(qr_buffer, 'PNG')
    qr_buffer.seek(0)
    
    qr_temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f'qr_{doc_number}.png')
    with open(qr_temp_path, 'wb') as f:
        f.write(qr_buffer.read())
    
    # Добавляем QR-код в PDF
    c.drawImage(qr_temp_path, width - 150, 50, width=100, height=100)
    
    # PIN-код рядом с QR-кодом
    pin_code = document_data.get('pin_code', '')
    c.setFont("Helvetica-Bold", 24)
    c.drawCentredString(width - 100, 30, pin_code)
    
    # DMED логотип (текстовый)
    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(colors.HexColor("#2563EB"))
    c.drawString(50, 80, "✱ DMED")
    c.setFont("Helvetica", 8)
    c.setFillColor(colors.black)
    c.drawString(50, 65, "Documents")
    
    # Информация внизу
    c.setFont("Helvetica", 7)
    c.drawString(50, 40, f"Hujjat DMED Yugoria tibbiy anborot tizimida hujalning'o'z kodini kiritish, yokl QR-kod orqali")
    c.drawString(50, 30, "tashabukor muman")
    
    # UUID документа
    document_uuid = document_data.get('uuid', '')
    if document_uuid:
        c.drawString(50, 20, f"Hujjat ID: {document_uuid}")
    else:
        c.drawString(50, 20, f"Hujjat ID: {doc_number}")
    c.drawString(50, 10, f"Yaratish sanasi: {issue_date_str}")
    
    c.save()
    
    # Удаляем временный QR-код
    if os.path.exists(qr_temp_path):
        os.remove(qr_temp_path)
    
    # Читаем PDF файл и сохраняем в хранилище (MinIO или локально)
    with open(filepath, 'rb') as f:
        pdf_data = f.read()
    
    # Сохраняем в хранилище
    stored_path = storage_manager.save_file(pdf_data, filename, 'application/pdf')
    
    # Если используется MinIO, удаляем локальный файл
    if storage_manager.use_minio and os.path.exists(filepath):
        try:
            os.remove(filepath)
        except:
            pass
    
    # Возвращаем путь (для MinIO это object_name, для локального - filepath)
    return stored_path


def fill_docx_template(document_data, template_path=None):
    """Заполняет DOCX шаблон данными из формы"""
    try:
        # Определяем путь к шаблону
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            # Используем стандартный шаблон или создаем новый
            default_template = os.path.join(app.config['TEMPLATE_FOLDER'], 'template.docx')
            if os.path.exists(default_template):
                doc = Document(default_template)
            else:
                # Создаем новый документ если шаблона нет
                doc = create_default_docx_template()
        
        # Подготавливаем данные для замены
        replacements = {
            '{{doc_number}}': document_data.get('doc_number', ''),
            '{{pin_code}}': document_data.get('pin_code', ''),
            '{{uuid}}': document_data.get('uuid', ''),
            '{{patient_name}}': document_data.get('patient_name', ''),
            '{{gender}}': document_data.get('gender', ''),
            '{{age}}': document_data.get('age', ''),
            '{{jshshir}}': document_data.get('jshshir', ''),
            '{{address}}': document_data.get('address', ''),
            '{{attached_medical_institution}}': document_data.get('attached_medical_institution', ''),
            '{{diagnosis}}': document_data.get('diagnosis', ''),
            '{{diagnosis_icd10_code}}': document_data.get('diagnosis_icd10_code', ''),
            '{{final_diagnosis}}': document_data.get('final_diagnosis', ''),
            '{{final_diagnosis_icd10_code}}': document_data.get('final_diagnosis_icd10_code', ''),
            '{{organization}}': document_data.get('organization', ''),
            '{{doctor_name}}': document_data.get('doctor_name', ''),
            '{{doctor_position}}': document_data.get('doctor_position', ''),
            '{{department_head_name}}': document_data.get('department_head_name', ''),
        }
        
        # Форматируем даты освобождения
        days_off_from = document_data.get('days_off_from')
        days_off_to = document_data.get('days_off_to')
        
        if days_off_from:
            if isinstance(days_off_from, str):
                try:
                    dt = datetime.fromisoformat(days_off_from.replace('Z', '+00:00'))
                    replacements['{{days_off_from}}'] = dt.strftime("%d.%m.%Y")
                except:
                    replacements['{{days_off_from}}'] = days_off_from[:10] if len(days_off_from) >= 10 else days_off_from
            else:
                replacements['{{days_off_from}}'] = days_off_from.strftime("%d.%m.%Y")
        else:
            replacements['{{days_off_from}}'] = ''
        
        if days_off_to:
            if isinstance(days_off_to, str):
                try:
                    dt = datetime.fromisoformat(days_off_to.replace('Z', '+00:00'))
                    replacements['{{days_off_to}}'] = dt.strftime("%d.%m.%Y")
                except:
                    replacements['{{days_off_to}}'] = days_off_to[:10] if len(days_off_to) >= 10 else days_off_to
            else:
                replacements['{{days_off_to}}'] = days_off_to.strftime("%d.%m.%Y")
        else:
            replacements['{{days_off_to}}'] = ''
        
        # Формируем строку периода освобождения
        if replacements['{{days_off_from}}'] and replacements['{{days_off_to}}']:
            replacements['{{days_off_period}}'] = f"{replacements['{{days_off_from}}']} - {replacements['{{days_off_to}}']}"
        elif replacements['{{days_off_from}}']:
            replacements['{{days_off_period}}'] = replacements['{{days_off_from}}']
        else:
            replacements['{{days_off_period}}'] = ''
        
        # Форматируем дату выдачи документа (с временем)
        issue_date = document_data.get('issue_date')
        if issue_date:
            if isinstance(issue_date, str):
                try:
                    dt = datetime.fromisoformat(issue_date.replace('Z', '+00:00'))
                    date_str = dt.strftime("%d.%m.%Y %H:%M")
                except:
                    # Если не удалось распарсить, пробуем другой формат
                    try:
                        dt = datetime.strptime(issue_date[:16], "%Y-%m-%dT%H:%M")
                        date_str = dt.strftime("%d.%m.%Y %H:%M")
                    except:
                        date_str = issue_date[:10] if len(issue_date) >= 10 else issue_date
            else:
                date_str = issue_date.strftime("%d.%m.%Y %H:%M")
        else:
            date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        
        replacements['{{issue_date}}'] = date_str
        replacements['{{current_date}}'] = datetime.now().strftime("%d.%m.%Y %H:%M")
        
        # Флаг для отслеживания первого вхождения {{pin_code}}
        pin_code_first_occurrence = {'found': False}
        
        # Функция для замены плейсхолдеров с сохранением форматирования и установкой Times New Roman
        def replace_placeholders_with_font(paragraph, replacements, font_name=DOCX_FONT_NAME):
            """Заменяет плейсхолдеры в параграфе, устанавливая шрифт Times New Roman для замененных значений.
            Только первое вхождение {{pin_code}} получает размер 24pt, остальные - обычный размер."""
            try:
                full_text = paragraph.text
                
                # Проверяем, есть ли плейсхолдеры
                has_placeholder = False
                for key in replacements.keys():
                    if key in full_text:
                        has_placeholder = True
                        break
                
                if not has_placeholder:
                    return
                
                # Сохраняем выравнивание и стиль параграфа
                alignment = paragraph.alignment
                
                # Сохраняем размер шрифта из первого run (если есть)
                original_font_size = None
                if paragraph.runs:
                    try:
                        original_font_size = paragraph.runs[0].font.size
                    except:
                        pass
                
                # Очищаем параграф
                paragraph.clear()
                if alignment:
                    paragraph.alignment = alignment
                
                # Разбиваем текст на части и обрабатываем каждый плейсхолдер отдельно
                import re
                # Создаем паттерн для поиска всех плейсхолдеров
                pattern = '|'.join(re.escape(key) for key in replacements.keys())
                parts = re.split(f'({pattern})', full_text)
                
                for part in parts:
                    if not part:
                        continue
                    
                    # Проверяем, является ли часть плейсхолдером
                    if part in replacements:
                        value = str(replacements[part])
                        run = paragraph.add_run(value)
                        
                        # Специальная обработка для {{pin_code}} - первое вхождение обычный размер, второе и последующие - 24pt
                        if part == '{{pin_code}}':
                            if not pin_code_first_occurrence['found']:
                                # Первое вхождение - обычный размер
                                try:
                                    run.font.name = font_name
                                except:
                                    try:
                                        run.font.name = 'Times New Roman'
                                    except:
                                        pass
                                
                                # Восстанавливаем размер шрифта
                                if original_font_size:
                                    try:
                                        run.font.size = original_font_size
                                    except:
                                        pass
                                pin_code_first_occurrence['found'] = True
                            else:
                                # Второе и последующие вхождения - размер 24pt
                                run.font.size = Pt(24)
                                run.font.name = 'Calibri'  # Используем Calibri для PIN-кода
                                run.bold = True
                        else:
                            # Для остальных плейсхолдеров используем стандартный шрифт
                            try:
                                run.font.name = font_name
                            except:
                                try:
                                    run.font.name = 'Times New Roman'
                                except:
                                    pass
                            
                            # Восстанавливаем размер шрифта
                            if original_font_size:
                                try:
                                    run.font.size = original_font_size
                                except:
                                    pass
                    else:
                        # Обычный текст - сохраняем оригинальное форматирование
                        run = paragraph.add_run(part)
                        try:
                            run.font.name = font_name
                        except:
                            try:
                                run.font.name = 'Times New Roman'
                            except:
                                pass
                        if original_font_size:
                            try:
                                run.font.size = original_font_size
                            except:
                                pass
            except Exception as e:
                print(f"Ошибка при замене плейсхолдеров в параграфе: {e}")
                import traceback
                print(traceback.format_exc())
                # В случае ошибки просто заменяем текст без форматирования
                try:
                    # Пробуем заменить плейсхолдеры простым способом
                    simple_text = full_text
                    for key, value in replacements.items():
                        simple_text = simple_text.replace(key, str(value))
                    paragraph.text = simple_text
                except Exception as fallback_error:
                    print(f"Ошибка при fallback замене: {fallback_error}")
        
        # Заменяем плейсхолдеры в параграфах
        for paragraph in doc.paragraphs:
            replace_placeholders_with_font(paragraph, replacements)
        
        # Заменяем плейсхолдеры в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_with_font(paragraph, replacements)
        
        # Заменяем плейсхолдеры в колонтитулах (headers и footers)
        for section in doc.sections:
            # Обрабатываем верхний колонтитул (header)
            if section.header:
                for paragraph in section.header.paragraphs:
                    replace_placeholders_with_font(paragraph, replacements)
                # Обрабатываем таблицы в header
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_placeholders_with_font(paragraph, replacements)
            
            # Обрабатываем нижний колонтитул (footer)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    replace_placeholders_with_font(paragraph, replacements)
                # Обрабатываем таблицы в footer
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_placeholders_with_font(paragraph, replacements)
        
        # Добавляем QR-код
        add_qr_code_to_docx(doc, document_data.get('pin_code', ''))
        
        # Сохраняем заполненный документ
        upload_folder = app.config.get('UPLOAD_FOLDER', 'static/generated_documents')
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder, exist_ok=True)
            print(f"Создана папка для документов: {upload_folder}")
        
        # Используем UUID для имени файла
        document_uuid = document_data.get('uuid', '')
        if not document_uuid:
            # Если UUID нет, генерируем новый
            document_uuid = str(uuid.uuid4())
        output_path = os.path.join(upload_folder, f"{document_uuid}.docx")
        
        print(f"Сохраняем документ: {output_path}")
        try:
            doc.save(output_path)
            print(f"Документ успешно сохранен: {output_path}")
            
            if os.path.exists(output_path):
                # Читаем DOCX файл и сохраняем в хранилище (MinIO или локально)
                with open(output_path, 'rb') as f:
                    docx_data = f.read()
                
                # Сохраняем в хранилище
                docx_filename = f"{document_uuid}.docx"
                stored_path = storage_manager.save_file(docx_data, docx_filename, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                
                # Если используется MinIO, удаляем локальный файл
                if storage_manager.use_minio and os.path.exists(output_path):
                    try:
                        os.remove(output_path)
                    except:
                        pass
                
                # Возвращаем путь (для MinIO это object_name, для локального - filepath)
                return stored_path
            else:
                print(f"ОШИБКА: Файл не был создан: {output_path}")
                return None
        except Exception as save_error:
            print(f"ОШИБКА при сохранении документа: {save_error}")
            import traceback
            print(traceback.format_exc())
            return None
        
    except Exception as e:
        import traceback
        error_msg = f"Ошибка при заполнении DOCX шаблона: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        app.logger.error(error_msg)
        return None


def create_default_docx_template():
    """Создает стандартный DOCX шаблон если его нет"""
    doc = Document()
    
    # Заголовок
    try:
        title = doc.add_heading("Ta'lim olayotgan shaxslar uchun mehnatga layoqatsizlik ma'lumotnomasi", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        # Если стиль Heading не доступен, создаем обычный параграф
        title = doc.add_paragraph("Ta'lim olayotgan shaxslar uchun mehnatga layoqatsizlik ma'lumotnomasi")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].font.bold = True
            title.runs[0].font.name = DOCX_FONT_NAME
    
    # Организация
    org_para = doc.add_paragraph("O'zbekiston Respublikasi Sog'liqni saqlash vazirligi")
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_para = doc.add_paragraph("4 - sonli oilaviy poliklinika")
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Номер документа
    doc.add_paragraph(f"Ro'yhatga olingan sana: {{issue_date}}")
    doc.add_paragraph(f"No {{doc_number}}")
    
    # Данные пациента
    try:
        doc.add_heading('Vaqtincha mehnatga layoqatsiz fuqaro haqidagi ma\'lumotlar:', 1)
    except:
        # Если стиль Heading 1 не доступен, создаем обычный параграф с жирным текстом
        para = doc.add_paragraph('Vaqtincha mehnatga layoqatsiz fuqaro haqidagi ma\'lumotlar:')
        para.runs[0].font.bold = True
        para.runs[0].font.name = DOCX_FONT_NAME
    doc.add_paragraph(f"FISH: {{patient_name}}")
    doc.add_paragraph(f"Jinsi: {{gender}}")
    doc.add_paragraph(f"JShShIR: {{jshshir}}")
    doc.add_paragraph(f"Yoshi: {{age}}")
    doc.add_paragraph(f"Yashash manzili: {{address}}")
    doc.add_paragraph(f"Biriktirilgan tibbiy muassasa: {{attached_medical_institution}}")
    doc.add_paragraph(f"Tashxis (KXT-10 kodi va Nomi): {{diagnosis_icd10_code}}: {{diagnosis}}")
    doc.add_paragraph(f"Yakuniy tashxis (Nomi va KXT-10 kodi): {{final_diagnosis_icd10_code}}: {{final_diagnosis}}")
    doc.add_paragraph(f"Ro'yhatga olingan sana: {{issue_date}}")
    doc.add_paragraph(f"No: {{doc_number}}")
    doc.add_paragraph(f"Davolovchi shifokor FISH: {{doctor_name}}")
    doc.add_paragraph(f"Bo'lim boshlig'i (mas'ul shaxs) FISH: {{department_head_name}}")
    doc.add_paragraph(f"Ishdan ozod etilgan kunlar: {{days_off_period}}")
    
    # Данные врача
    try:
        doc.add_heading('Shifokor ma\'lumotlari:', 1)
    except:
        # Если стиль Heading 1 не доступен, создаем обычный параграф с жирным текстом
        para = doc.add_paragraph('Shifokor ma\'lumotlari:')
        if para.runs:
            para.runs[0].font.bold = True
            para.runs[0].font.name = DOCX_FONT_NAME
        else:
            run = para.add_run('Shifokor ma\'lumotlari:')
            run.font.bold = True
            run.font.name = DOCX_FONT_NAME
    doc.add_paragraph(f"FISH: {{doctor_name}}")
    doc.add_paragraph(f"Lavozim: {{doctor_position}}")
    
    # PIN-код
    doc.add_paragraph(f"PIN-kod: {{pin_code}}")
    
    return doc


def add_qr_code_to_docx(doc, pin_code, replacements=None):
    """Добавляет QR-код в DOCX документ, заменяя плейсхолдер {{qr_code}} если он есть"""
    try:
        # Генерируем QR-код
        try:
            qr_url = url_for('verify_document', _external=True)
        except RuntimeError:
            # Если url_for не работает вне контекста запроса, используем относительный путь
            qr_url = '/verify'
        
        qr_img = generate_qr_code(qr_url)
        
        # Сохраняем QR-код во временный файл
        qr_temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f'qr_temp_{pin_code}.png')
        if not os.path.exists(os.path.dirname(qr_temp_path)):
            os.makedirs(os.path.dirname(qr_temp_path), exist_ok=True)
        # PIL Image.save() принимает путь к файлу (строку) или файловый объект
        # Используем файловый объект для совместимости с типизацией
        with open(qr_temp_path, 'wb') as f:
            qr_img.save(f, 'PNG')
        
        qr_placeholder_found = False
        
        # Функция для замены плейсхолдера {{qr_code}} на изображение
        def replace_qr_placeholder(paragraph):
            nonlocal qr_placeholder_found
            if '{{qr_code}}' in paragraph.text:
                qr_placeholder_found = True
                # Сохраняем выравнивание
                alignment = paragraph.alignment
                # Очищаем параграф
                paragraph.clear()
                if alignment:
                    paragraph.alignment = alignment
                # Добавляем QR-код изображение
                run = paragraph.add_run()
                run.add_picture(qr_temp_path, width=Inches(1.5))
                return True
            return False
        
        # Ищем плейсхолдер в основном теле документа
        for paragraph in doc.paragraphs:
            if replace_qr_placeholder(paragraph):
                break
        
        # Ищем плейсхолдер в таблицах основного тела
        if not qr_placeholder_found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if replace_qr_placeholder(paragraph):
                                qr_placeholder_found = True
                                break
                        if qr_placeholder_found:
                            break
                    if qr_placeholder_found:
                        break
                if qr_placeholder_found:
                    break
        
        # Ищем плейсхолдер в колонтитулах (headers и footers)
        if not qr_placeholder_found:
            for section in doc.sections:
                # Проверяем header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if replace_qr_placeholder(paragraph):
                            qr_placeholder_found = True
                            break
                    if not qr_placeholder_found:
                        # Проверяем таблицы в header
                        for table in section.header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if replace_qr_placeholder(paragraph):
                                            qr_placeholder_found = True
                                            break
                                    if qr_placeholder_found:
                                        break
                                if qr_placeholder_found:
                                    break
                            if qr_placeholder_found:
                                break
                
                if qr_placeholder_found:
                    break
                
                # Проверяем footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if replace_qr_placeholder(paragraph):
                            qr_placeholder_found = True
                            break
                    if not qr_placeholder_found:
                        # Проверяем таблицы в footer
                        for table in section.footer.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if replace_qr_placeholder(paragraph):
                                            qr_placeholder_found = True
                                            break
                                    if qr_placeholder_found:
                                        break
                                if qr_placeholder_found:
                                    break
                            if qr_placeholder_found:
                                break
                
                if qr_placeholder_found:
                    break
        
        # Если плейсхолдер не найден, добавляем QR-код в конец документа
        if not qr_placeholder_found:
            last_para = doc.add_paragraph()
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = last_para.add_run()
            run.add_picture(qr_temp_path, width=Inches(1.5))
        
        # Удаляем временный файл
        if os.path.exists(qr_temp_path):
            os.remove(qr_temp_path)
            
    except Exception as e:
        print(f"Ошибка при добавлении QR-кода: {e}")
        import traceback
        print(traceback.format_exc())


# Маршруты

@app.route('/')
def index():
    """Главная страница с формой"""
    return render_template('index.html')


@app.route('/create-document', methods=['POST'])
def create_document():
    """Создание нового документа"""
    try:
        data = request.get_json()
        
        # Генерируем уникальные идентификаторы
        doc_number = generate_document_number()
        pin_code = generate_pin_code()
        document_uuid = str(uuid.uuid4())
        
        # Подготавливаем данные для вставки
        # Обрабатываем даты освобождения
        days_off_from = data.get('days_off_from')
        days_off_to = data.get('days_off_to')
        
        # Обрабатываем дату выдачи документа
        issue_date_input = data.get('issue_date')
        if issue_date_input:
            # Если дата указана в форме, добавляем текущее время
            # Формат из формы: YYYY-MM-DD, добавляем время HH:MM:SS
            try:
                date_part = datetime.strptime(issue_date_input, "%Y-%m-%d")
                # Добавляем текущее время
                now = datetime.now()
                issue_date = date_part.replace(hour=now.hour, minute=now.minute, second=now.second).isoformat()
            except:
                # Если не удалось распарсить, используем как есть
                issue_date = issue_date_input
        else:
            # Если не указана, используем текущую дату и время
            issue_date = datetime.now().isoformat()
        
        document_data = {
            'doc_number': doc_number,
            'pin_code': pin_code,
            'uuid': document_uuid,
            'patient_name': data.get('patient_name'),
            'gender': data.get('gender'),
            'age': data.get('age'),
            'jshshir': data.get('jshshir'),
            'address': data.get('address'),
            'attached_medical_institution': data.get('attached_medical_institution'),
            'diagnosis': data.get('diagnosis'),
            'diagnosis_icd10_code': data.get('diagnosis_icd10_code'),
            'final_diagnosis': data.get('final_diagnosis'),
            'final_diagnosis_icd10_code': data.get('final_diagnosis_icd10_code'),
            'organization': data.get('organization'),
            'doctor_name': data.get('doctor_name'),
            'doctor_position': data.get('doctor_position'),
            'department_head_name': data.get('department_head_name'),
            'days_off_from': days_off_from if days_off_from else None,
            'days_off_to': days_off_to if days_off_to else None,
            'issue_date': issue_date
        }
        
        # Вставляем документ в БД
        try:
            print(f"DEBUG: Вставка документа в БД...")
            print(f"DEBUG: doc_number = {document_data.get('doc_number')}")
            print(f"DEBUG: pin_code = {document_data.get('pin_code')}")
            print(f"DEBUG: patient_name = {document_data.get('patient_name')}")
            
            created_document = db_insert('documents', document_data)
            
            if not created_document:
                print("ERROR: db_insert вернул None")
                return jsonify({'success': False, 'error': 'Ошибка создания документа'}), 500
            
            document_id = created_document['id']
            print(f"DEBUG: Документ создан с ID: {document_id}")
        except Exception as db_error:
            print(f"ERROR: Ошибка при создании документа: {db_error}")
            import traceback
            print(traceback.format_exc())
            return jsonify({'success': False, 'error': f'Ошибка создания документа: {str(db_error)}'}), 500
        
        # Генерируем DOCX документ из шаблона (временный, для конвертации)
        docx_path = fill_docx_template(created_document)
        
        if not docx_path:
            error_detail = "Не удалось создать DOCX документ. Проверьте логи сервера."
            print(f"ОШИБКА: {error_detail}")
            import traceback
            print(traceback.format_exc())
            return jsonify({'success': False, 'error': error_detail}), 500
        
        # Конвертируем DOCX в PDF сразу после создания
        pdf_path = convert_docx_to_pdf_from_docx(docx_path, created_document)
        
        # Удаляем временный DOCX файл после конвертации
        if docx_path and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
                print(f"[OK] Временный DOCX файл удален: {docx_path}")
            except Exception as e:
                print(f"[WARNING] Не удалось удалить временный DOCX файл: {e}")
        
        # Если конвертация не удалась, используем базовый метод генерации PDF
        if not pdf_path or not storage_manager.file_exists(pdf_path):
            print("[WARNING] Конвертация DOCX->PDF не удалась, используем базовый метод")
            pdf_path = generate_pdf_document(created_document)
        
        if not pdf_path:
            error_detail = "Не удалось создать PDF документ. Проверьте логи сервера."
            print(f"ОШИБКА: {error_detail}")
            return jsonify({'success': False, 'error': error_detail}), 500
        
        # Обновляем путь к PDF в базе данных
        try:
            db_update('documents', {'pdf_path': pdf_path}, 'id = %s', [document_id])
        except Exception as db_error:
            print(f"Предупреждение: Не удалось обновить путь к PDF в БД: {db_error}")
        
        return jsonify({
            'success': True,
            'document_id': document_id,
            'doc_number': doc_number,
            'pin_code': pin_code,
            'download_url': url_for('download_document', doc_id=document_id)
        })
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/download/<int:doc_id>')
def download_document(doc_id):
    """Скачивание PDF документа"""
    try:
        document = db_select('documents', 'id = %s', [doc_id], fetch_one=True)
        
        if not document:
            return "Документ не найден", 404
        
        document_uuid = document.get('uuid', '')
        if not document_uuid:
            document_uuid = str(uuid.uuid4())
        
        pdf_path = document.get('pdf_path')
        filename = f'{document_uuid}.pdf'
        
        # Пытаемся получить файл из хранилища
        if pdf_path:
            file_data = storage_manager.get_file(pdf_path)
            if file_data:
                # Возвращаем файл из хранилища
                return send_file(
                    BytesIO(file_data),
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/pdf'
                )
        
        # Если файл не найден в хранилище, проверяем локально (для обратной совместимости)
        if pdf_path and os.path.exists(pdf_path):
            return send_file(
                pdf_path,
                as_attachment=True,
                download_name=filename
            )
        
        # Если PDF не существует, генерируем его заново
        pdf_path = generate_pdf_document(document)
        if pdf_path:
            # Обновляем путь в БД
            db_update('documents', {'pdf_path': pdf_path}, 'id = %s', [doc_id])
            # Получаем файл из хранилища
            file_data = storage_manager.get_file(pdf_path)
            if file_data:
                return send_file(
                    BytesIO(file_data),
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/pdf'
                )
            # Fallback на локальный файл
            if os.path.exists(pdf_path):
                return send_file(
                    pdf_path,
                    as_attachment=True,
                    download_name=filename
                )
        return "PDF файл не найден", 404
    except Exception as e:
        return f"Ошибка: {str(e)}", 500


@app.route('/download-by-uuid/<uuid>')
def download_by_uuid(uuid):
    """Скачивание PDF документа по UUID (для QR-кода)"""
    try:
        document = db_select('documents', 'uuid = %s', [uuid], fetch_one=True)
        
        if not document:
            return "Документ не найден", 404
        
        pdf_path = document.get('pdf_path')
        filename = f'{uuid}.pdf'
        
        # Пытаемся получить файл из хранилища
        if pdf_path:
            file_data = storage_manager.get_file(pdf_path)
            if file_data:
                # Возвращаем файл из хранилища
                return send_file(
                    BytesIO(file_data),
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/pdf'
                )
        
        # Если файл не найден в хранилище, проверяем локально (для обратной совместимости)
        if pdf_path and os.path.exists(pdf_path):
            return send_file(
                pdf_path,
                as_attachment=True,
                download_name=filename
            )
        
        # Если PDF не существует, генерируем его заново
        pdf_path = generate_pdf_document(document)
        if pdf_path:
            # Обновляем путь в БД
            db_update('documents', {'pdf_path': pdf_path}, 'uuid = %s', [uuid])
            # Получаем файл из хранилища
            file_data = storage_manager.get_file(pdf_path)
            if file_data:
                return send_file(
                    BytesIO(file_data),
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/pdf'
                )
            # Fallback на локальный файл
            if os.path.exists(pdf_path):
                return send_file(
                    pdf_path,
                    as_attachment=True,
                    download_name=filename
                )
        return "PDF файл не найден", 404
    except Exception as e:
        return f"Ошибка: {str(e)}", 500


@app.route('/download-docx/<int:doc_id>')
def download_docx(doc_id):
    """Скачивание DOCX документа"""
    try:
        document = db_select('documents', 'id = %s', [doc_id], fetch_one=True)
        
        if not document:
            return "Документ не найден", 404
        
        docx_path = document.get('docx_path')
        doc_number = document.get('doc_number', 'unknown')
        filename = f'spravka_{doc_number}.docx'
        
        # Пытаемся получить файл из хранилища
        if docx_path:
            file_data = storage_manager.get_file(docx_path)
            if file_data:
                return send_file(
                    BytesIO(file_data),
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
        
        # Если файл не найден в хранилище, проверяем локально (для обратной совместимости)
        if docx_path and os.path.exists(docx_path):
            return send_file(
                docx_path,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        
        # Если DOCX не существует, генерируем его заново
        docx_path = fill_docx_template(document)
        if docx_path:
            # Обновляем путь в БД
            db_update('documents', {'docx_path': docx_path}, 'id = %s', [doc_id])
            # Получаем файл из хранилища
            file_data = storage_manager.get_file(docx_path)
            if file_data:
                return send_file(
                    BytesIO(file_data),
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            # Fallback на локальный файл
            if os.path.exists(docx_path):
                return send_file(
                    docx_path,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
        return "DOCX файл не найден", 404
    except Exception as e:
        return f"Ошибка: {str(e)}", 500


def convert_docx_to_pdf_from_docx(docx_path, document_data, output_path=None):
    """Конвертирует DOCX файл в PDF, читая содержимое DOCX и создавая PDF"""
    # Получаем DOCX файл из хранилища или локально
    docx_data = None
    temp_docx_path = None
    
    if docx_path:
        # Пытаемся получить из хранилища
        docx_data = storage_manager.get_file(docx_path)
        if docx_data:
            # Сохраняем во временный файл для конвертации
            temp_docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f'temp_{uuid.uuid4()}.docx')
            with open(temp_docx_path, 'wb') as f:
                f.write(docx_data)
            docx_path = temp_docx_path
    
    if not docx_path or not os.path.exists(docx_path):
        return None
    
    try:
        if output_path is None:
            # Используем UUID для имени файла
            document_uuid = document_data.get('uuid', '') if document_data else ''
            if not document_uuid:
                document_uuid = str(uuid.uuid4())
            output_path = os.path.join(
                app.config['UPLOAD_FOLDER'],
                f"{document_uuid}.pdf"
            )
        
        # Метод 1: Используем mammoth + weasyprint (не требует LibreOffice/Word, сохраняет форматирование)
        if MAMMOTH_AVAILABLE and WEASYPRINT_AVAILABLE:
            try:
                # Функция для конвертации изображений в base64
                def convert_image(image):
                    """Конвертирует изображения из DOCX в base64 для вставки в HTML"""
                    with image.open() as image_bytes:
                        image_base64 = base64.b64encode(image_bytes.read()).decode("utf-8")
                        return {"src": f"data:{image.content_type};base64,{image_base64}"}
                
                # Конвертируем DOCX в HTML через mammoth с обработкой изображений
                with open(docx_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(
                        docx_file,
                        convert_image=mammoth.images.img_element(convert_image)
                    )
                    html_content = result.value
                    # Получаем предупреждения (если есть)
                    if result.messages:
                        print(f"[WARNING] Предупреждения mammoth: {result.messages}")
                
                # Добавляем базовые стили для лучшего отображения
                html_with_styles = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        @page {{
                            size: A4;
                            margin: 2cm;
                        }}
                        body {{
                            font-family: 'Times New Roman', serif;
                            font-size: 12pt;
                            line-height: 1.5;
                            color: #000;
                        }}
                        p {{
                            margin: 0.5em 0;
                        }}
                        table {{
                            border-collapse: collapse;
                            width: 100%;
                            margin: 1em 0;
                        }}
                        table td, table th {{
                            border: 1px solid #ddd;
                            padding: 8px;
                        }}
                        table th {{
                            background-color: #f2f2f2;
                            font-weight: bold;
                        }}
                        img {{
                            max-width: 100%;
                            height: auto;
                        }}
                        h1, h2, h3, h4, h5, h6 {{
                            margin: 1em 0 0.5em 0;
                            font-weight: bold;
                        }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                
                # Конвертируем HTML в PDF через weasyprint
                HTML(string=html_with_styles).write_pdf(output_path)
                
                if os.path.exists(output_path):
                    print(f"[OK] DOCX успешно конвертирован в PDF через mammoth+weasyprint: {output_path}")
                    # Сохраняем PDF в хранилище
                    with open(output_path, 'rb') as f:
                        pdf_data = f.read()
                    
                    pdf_filename = os.path.basename(output_path)
                    stored_path = storage_manager.save_file(pdf_data, pdf_filename, 'application/pdf')
                    
                    # Удаляем локальный файл если используется MinIO
                    if storage_manager.use_minio and os.path.exists(output_path):
                        try:
                            os.remove(output_path)
                        except:
                            pass
                    
                    # Удаляем временный DOCX файл если был создан
                    if temp_docx_path and os.path.exists(temp_docx_path):
                        try:
                            os.remove(temp_docx_path)
                        except:
                            pass
                    
                    return stored_path
            except Exception as e:
                print(f"[WARNING] Ошибка при конвертации через mammoth+weasyprint: {e}")
                import traceback
                print(traceback.format_exc())
                print("Пробуем альтернативный метод...")
        
        # Метод 2: Попытка использовать docx2pdf (требует LibreOffice/Word)
        if DOCX2PDF_AVAILABLE:
            try:
                convert(docx_path, output_path)
                
                if os.path.exists(output_path):
                    print(f"[OK] DOCX успешно конвертирован в PDF через docx2pdf: {output_path}")
                    # Сохраняем PDF в хранилище
                    with open(output_path, 'rb') as f:
                        pdf_data = f.read()
                    
                    pdf_filename = os.path.basename(output_path)
                    stored_path = storage_manager.save_file(pdf_data, pdf_filename, 'application/pdf')
                    
                    # Удаляем локальный файл если используется MinIO
                    if storage_manager.use_minio and os.path.exists(output_path):
                        try:
                            os.remove(output_path)
                        except:
                            pass
                    
                    # Удаляем временный DOCX файл если был создан
                    if temp_docx_path and os.path.exists(temp_docx_path):
                        try:
                            os.remove(temp_docx_path)
                        except:
                            pass
                    
                    return stored_path
            except Exception as e:
                print(f"[WARNING] Ошибка при конвертации через docx2pdf: {e}")
                print("Используем альтернативный метод...")
        
        # Метод 3: Читаем DOCX и создаем PDF на основе его содержимого (базовый метод)
        try:
            doc = Document(docx_path)
            
            if output_path is None:
                output_path = docx_path.replace('.docx', '.pdf')
            
            # Создаем PDF
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            y_position = height - 50
            line_height = 20
            margin = 50
            
            # Извлекаем текст из DOCX
            for paragraph in doc.paragraphs:
                if y_position < 100:  # Новая страница если нужно
                    c.showPage()
                    y_position = height - 50
                
                text = paragraph.text.strip()
                if text:
                    # Определяем стиль текста
                    is_bold = False
                    font_size = 10
                    
                    for run in paragraph.runs:
                        if run.bold:
                            is_bold = True
                        if run.font.size:
                            font_size = run.font.size.pt if run.font.size else 10
                    
                    # Устанавливаем шрифт
                    font_name = "Helvetica-Bold" if is_bold else "Helvetica"
                    c.setFont(font_name, font_size)
                    
                    # Разбиваем длинный текст на строки
                    max_width = width - 2 * margin
                    words = text.split()
                    lines = []
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        if c.stringWidth(test_line, font_name, font_size) <= max_width:
                            current_line = test_line
                        else:
                            if current_line:
                                lines.append(current_line)
                            current_line = word
                    if current_line:
                        lines.append(current_line)
                    
                    # Рисуем строки
                    for line in lines:
                        if y_position < 100:
                            c.showPage()
                            y_position = height - 50
                        c.drawString(margin, y_position, line)
                        y_position -= line_height
            
            # Обрабатываем таблицы
            for table in doc.tables:
                if y_position < 150:
                    c.showPage()
                    y_position = height - 50
                
                y_position -= 20  # Отступ перед таблицей
                c.setFont("Helvetica-Bold", 10)
                
                for row in table.rows:
                    if y_position < 100:
                        c.showPage()
                        y_position = height - 50
                    
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        # Обрезаем длинный текст
                        if len(row_text) > 100:
                            row_text = row_text[:97] + "..."
                        c.setFont("Helvetica", 9)
                        c.drawString(margin, y_position, row_text)
                        y_position -= line_height
            
            # Обрабатываем изображения (QR-код)
            # Ищем изображения в документе
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = doc.part.related_parts[rel.target_ref]
                        image_data = image_part.blob
                        
                        # Сохраняем изображение временно
                        img_temp = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_img.png')
                        with open(img_temp, 'wb') as f:
                            f.write(image_data)
                        
                        # Добавляем изображение в PDF
                        if y_position < 150:
                            c.showPage()
                            y_position = height - 50
                        
                        c.drawImage(img_temp, width - 150, 50, width=100, height=100)
                        
                        # Удаляем временный файл
                        if os.path.exists(img_temp):
                            os.remove(img_temp)
                    except:
                        pass
            
            c.save()
            
            if os.path.exists(output_path):
                print(f"[OK] DOCX успешно конвертирован в PDF (альтернативный метод): {output_path}")
                # Сохраняем PDF в хранилище
                with open(output_path, 'rb') as f:
                    pdf_data = f.read()
                
                pdf_filename = os.path.basename(output_path)
                stored_path = storage_manager.save_file(pdf_data, pdf_filename, 'application/pdf')
                
                # Удаляем локальный файл если используется MinIO
                if storage_manager.use_minio and os.path.exists(output_path):
                    try:
                        os.remove(output_path)
                    except:
                        pass
                
                # Удаляем временный DOCX файл если был создан
                if temp_docx_path and os.path.exists(temp_docx_path):
                    try:
                        os.remove(temp_docx_path)
                    except:
                        pass
                
                return stored_path
            
        except Exception as e:
            print(f"[WARNING] Ошибка при альтернативной конвертации: {e}")
            import traceback
            print(traceback.format_exc())
        
        # Метод 4: Fallback - генерируем PDF из данных БД (если есть)
        if document_data:
            print("Используем fallback: генерация PDF из данных БД")
            pdf_path = generate_pdf_document(document_data)
            if pdf_path:
                return pdf_path
        
        # Удаляем временный DOCX файл если был создан
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except:
                pass
        
        return None
        
    except Exception as e:
        print(f"[ERROR] Ошибка при конвертации DOCX в PDF: {e}")
        import traceback
        print(traceback.format_exc())
        # Удаляем временный DOCX файл если был создан
        if temp_docx_path and os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except:
                pass
        return None


@app.route('/convert-docx-to-pdf/<int:doc_id>')
def convert_docx_to_pdf_route(doc_id):
    """Конвертирует DOCX документ в PDF и возвращает PDF файл"""
    try:
        document = db_select('documents', 'id = %s', [doc_id], fetch_one=True)
        
        if not document:
            return jsonify({'success': False, 'error': 'Документ не найден'}), 404
        docx_path = document.get('docx_path')
        
        if not docx_path or not os.path.exists(docx_path):
            # Если DOCX нет, генерируем PDF из данных БД
            pdf_path = generate_pdf_document(document)
            if pdf_path and os.path.exists(pdf_path):
                document_uuid = document.get('uuid', '')
                if not document_uuid:
                    document_uuid = str(uuid.uuid4())
                return send_file(
                    pdf_path,
                    as_attachment=True,
                    download_name=f'{document_uuid}.pdf',
                    mimetype='application/pdf'
                )
            else:
                return jsonify({'success': False, 'error': 'Не удалось создать PDF'}), 500
        
        # Пытаемся конвертировать DOCX в PDF
        document_uuid = document.get('uuid', '')
        if not document_uuid:
            document_uuid = str(uuid.uuid4())
        pdf_output_path = os.path.join(
            app.config['UPLOAD_FOLDER'], 
            f"{document_uuid}.pdf"
        )
        
        converted_pdf = convert_docx_to_pdf_from_docx(docx_path, document, pdf_output_path)
        
        if converted_pdf and os.path.exists(converted_pdf):
            # Обновляем путь к PDF в БД
            try:
                db_update('documents', {'pdf_path': converted_pdf}, 'id = %s', [doc_id])
            except:
                pass  # Не критично, если не удалось обновить
            
            return send_file(
                converted_pdf,
                as_attachment=True,
                download_name=f'{document_uuid}.pdf',
                mimetype='application/pdf'
            )
        else:
            # Если конвертация не удалась, генерируем PDF из данных БД
            pdf_path = generate_pdf_document(document)
            if pdf_path and os.path.exists(pdf_path):
                return send_file(
                    pdf_path,
                    as_attachment=True,
                    download_name=f'{document_uuid}.pdf',
                    mimetype='application/pdf'
                )
            else:
                return jsonify({'success': False, 'error': 'Не удалось создать PDF из DOCX или из данных'}), 500
            
    except Exception as e:
        import traceback
        error_msg = f"Ошибка при конвертации: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/verify')
def verify_document():
    """Страница верификации документа"""
    return render_template('verify.html')


@app.route('/verify/<uuid>')
def verify_by_uuid(uuid):
    """Страница верификации документа по UUID (из QR-кода)"""
    try:
        # Ищем документ по UUID
        document = db_select('documents', 'uuid = %s', [uuid], fetch_one=True)
        
        if not document:
            return render_template('verify.html', error='Документ не найден')
        
        # Преобразуем в объект-подобный формат для шаблона
        class DocumentObj:
            def __init__(self, data):
                for key, value in data.items():
                    setattr(self, key, value)
            
            def __getitem__(self, key):
                return getattr(self, key, None)
        
        doc_obj = DocumentObj(document)
        return render_template('document_view.html', document=doc_obj, now=datetime.now)
    except Exception as e:
        return render_template('verify.html', error=f'Ошибка: {str(e)}')


@app.route('/verify-pin', methods=['POST'])
def verify_pin():
    """Проверка PIN-кода и возврат документа"""
    try:
        data = request.get_json()
        pin_code = data.get('pin_code')
        
        if not pin_code:
            return jsonify({'success': False, 'error': 'PIN-код не указан'}), 400
        
        # Ищем документ по PIN-коду в БД
        document = db_select('documents', 'pin_code = %s', [pin_code], fetch_one=True)
        
        if not document:
            return jsonify({'success': False, 'error': 'Документ не найден'}), 404
        
        # Форматируем дату
        issue_date = document.get('issue_date')
        if issue_date:
            if isinstance(issue_date, str):
                try:
                    dt = datetime.fromisoformat(issue_date.replace('Z', '+00:00'))
                    issue_date_str = dt.strftime("%d.%m.%Y %H:%M")
                except:
                    try:
                        dt = datetime.strptime(issue_date[:16], "%Y-%m-%dT%H:%M")
                        issue_date_str = dt.strftime("%d.%m.%Y %H:%M")
                    except:
                        issue_date_str = issue_date[:10] if len(issue_date) >= 10 else issue_date
            else:
                issue_date_str = issue_date.strftime("%d.%m.%Y %H:%M")
        else:
            issue_date_str = ''
        
        # Форматируем даты освобождения
        days_off_from = document.get('days_off_from')
        days_off_to = document.get('days_off_to')
        days_off_str = ''
        if days_off_from and days_off_to:
            try:
                from_dt = datetime.fromisoformat(days_off_from.replace('Z', '+00:00')) if isinstance(days_off_from, str) else days_off_from
                to_dt = datetime.fromisoformat(days_off_to.replace('Z', '+00:00')) if isinstance(days_off_to, str) else days_off_to
                days_off_str = f"{from_dt.strftime('%d.%m.%Y')} - {to_dt.strftime('%d.%m.%Y')}"
            except:
                days_off_str = f"{days_off_from} - {days_off_to}"
        elif days_off_from:
            try:
                from_dt = datetime.fromisoformat(days_off_from.replace('Z', '+00:00')) if isinstance(days_off_from, str) else days_off_from
                days_off_str = from_dt.strftime('%d.%m.%Y')
            except:
                days_off_str = str(days_off_from)
        
        # Возвращаем информацию о документе
        return jsonify({
            'success': True,
            'document': {
                'id': document.get('id'),
                'doc_number': document.get('doc_number'),
                'patient_name': document.get('patient_name'),
                'gender': document.get('gender'),
                'age': document.get('age'),
                'jshshir': document.get('jshshir'),
                'address': document.get('address'),
                'attached_medical_institution': document.get('attached_medical_institution'),
                'diagnosis': document.get('diagnosis'),
                'diagnosis_icd10_code': document.get('diagnosis_icd10_code'),
                'final_diagnosis': document.get('final_diagnosis'),
                'final_diagnosis_icd10_code': document.get('final_diagnosis_icd10_code'),
                'organization': document.get('organization'),
                'issue_date': issue_date_str,
                'doctor_name': document.get('doctor_name'),
                'doctor_position': document.get('doctor_position'),
                'department_head_name': document.get('department_head_name'),
                'days_off_period': days_off_str,
                'uuid': document.get('uuid', ''),
                'pdf_url': url_for('download_document', doc_id=document.get('id'), _external=True),
                'pdf_url_by_uuid': url_for('download_by_uuid', uuid=document.get('uuid', ''), _external=True) if document.get('uuid') else None
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/document/<int:doc_id>')
def view_document(doc_id):
    """Просмотр документа по ID"""
    try:
        document = db_select('documents', 'id = %s', [doc_id], fetch_one=True)
        
        if not document:
            return "Документ не найден", 404
        # Преобразуем в объект-подобный формат для шаблона
        class DocumentObj:
            def __init__(self, data):
                for key, value in data.items():
                    setattr(self, key, value)
            
            def __getitem__(self, key):
                return getattr(self, key, None)
        
        doc_obj = DocumentObj(document)
        return render_template('document_view.html', document=doc_obj, now=datetime.now)
    except Exception as e:
        return f"Ошибка: {str(e)}", 500


@app.route('/document/<string:doc_number>')
def view_document_by_number(doc_number):
    """Просмотр документа по номеру"""
    try:
        document = db_select('documents', 'doc_number = %s', [doc_number], fetch_one=True)
        
        if not document:
            return "Документ не найден", 404
        # Преобразуем в объект-подобный формат для шаблона
        class DocumentObj:
            def __init__(self, data):
                for key, value in data.items():
                    setattr(self, key, value)
            
            def __getitem__(self, key):
                return getattr(self, key, None)
        
        doc_obj = DocumentObj(document)
        return render_template('document_view.html', document=doc_obj, now=datetime.now)
    except Exception as e:
        return f"Ошибка: {str(e)}", 500


@app.route('/files')
def files_history():
    """Страница истории файлов"""
    return render_template('files_history.html')


@app.route('/api/files', methods=['GET'])
def list_files():
    """Получение списка всех файлов в хранилище с информацией из БД"""
    try:
        # Получаем параметры запроса
        prefix = request.args.get('prefix', '')
        file_type = request.args.get('type', '')  # 'pdf', 'docx' или пусто для всех
        
        # Формируем префикс для фильтрации
        search_prefix = prefix
        
        # Получаем список файлов из хранилища
        files = storage_manager.list_files(prefix=search_prefix, recursive=True)
        
        # Фильтруем по типу файла если указан
        if file_type:
            extension = f'.{file_type.lower()}'
            files = [f for f in files if f['name'].lower().endswith(extension)]
        
        # Обогащаем файлы информацией из БД
        enriched_files = []
        for file_info in files:
            # Извлекаем UUID из имени файла (убираем расширение)
            filename = file_info['name']
            uuid_from_filename = filename.replace('.pdf', '').replace('.docx', '')
            
            # Ищем документ в БД по UUID
            document = None
            try:
                document = db_select('documents', 'uuid = %s', [uuid_from_filename], fetch_one=True)
            except:
                pass
            
            # Добавляем информацию из БД
            enriched_file = {
                **file_info,
                'patient_name': document.get('patient_name') if document else None,
                'doc_number': document.get('doc_number') if document else None,
                'issue_date': document.get('issue_date').isoformat() if document and document.get('issue_date') else None,
            }
            
            enriched_files.append(enriched_file)
        
        # Форматируем даты для JSON
        for file_info in enriched_files:
            if isinstance(file_info.get('last_modified'), datetime):
                file_info['last_modified'] = file_info['last_modified'].isoformat()
        
        return jsonify({
            'success': True,
            'count': len(enriched_files),
            'files': enriched_files,
            'storage_type': 'minio' if storage_manager.use_minio else 'local'
        })
    except Exception as e:
        import traceback
        print(f"Ошибка при получении списка файлов: {e}")
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/files/download/<path:filename>', methods=['GET'])
def download_file_by_name(filename):
    """Скачивание файла по имени из хранилища"""
    try:
        # Получаем файл из хранилища
        file_data = storage_manager.get_file(filename)
        
        if not file_data:
            return jsonify({'success': False, 'error': 'Файл не найден'}), 404
        
        # Определяем MIME тип
        content_type = 'application/octet-stream'
        if filename.lower().endswith('.pdf'):
            content_type = 'application/pdf'
        elif filename.lower().endswith('.docx'):
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.lower().endswith('.doc'):
            content_type = 'application/msword'
        
        return send_file(
            BytesIO(file_data),
            as_attachment=True,
            download_name=filename,
            mimetype=content_type
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/files/delete/<path:filename>', methods=['DELETE'])
def delete_file_by_name(filename):
    """Удаление файла по имени из хранилища"""
    try:
        success = storage_manager.delete_file(filename)
        
        if success:
            return jsonify({'success': True, 'message': f'Файл {filename} успешно удален'})
        else:
            return jsonify({'success': False, 'error': 'Файл не найден или не удалось удалить'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':
    # Создаем директорию для документов
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    
    app.run(debug=True, host='0.0.0.0', port=5000)
